import streamlit as st
import requests
import time
import json
import os
import math
import base64
import asyncio
import concurrent.futures
import hashlib
import threading
import queue
import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime, timedelta
from collections import defaultdict
# PASSWORD PROTECTION
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.title("🔐 AQUALINE STUDIO")
    password = st.text_input("กรุณาใส่รหัสผ่าน", type="password")
    if st.button("เข้าสู่ระบบ"):
        if password == st.secrets.get("APP_PASSWORD", ""):
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("รหัสผ่านไม่ถูกต้อง")
    st.stop()

# ── PDF Export ──
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

# ==========================================
# 🔑 1. SECURITY & CONFIG 
# ==========================================
if "GOOGLE_API_KEY" in st.secrets:
    API_KEY = st.secrets["GOOGLE_API_KEY"]
else:
    st.error("🔑 ไม่พบรหัสใน secrets.toml (ต้องอยู่ในโฟลเดอร์ .streamlit)")
    st.stop()

# ── Chairman Master Config ──
ANTHROPIC_API_KEY    = st.secrets.get("ANTHROPIC_API_KEY", "")
CHAIRMAN_BUDGET_THB  = 175.0
CLAUDE_INPUT_USD     = 0.80
CLAUDE_OUTPUT_USD    = 4.00
USD_TO_THB           = float(st.secrets.get("USD_TO_THB", "35.0"))
CHAIRMAN_BUDGET_FILE = "chairman_budget.json"

VAULT_FILE = "project_vault.json"
ANALYTICS_FILE = "analytics_data.json"
CACHE_FILE = "brief_cache.json"

st.set_page_config(page_title="AQUALINE SPECIAL TEAM V9.0 ULTRA", layout="wide", initial_sidebar_state="expanded")

# ==========================================
# 🔴 FIX NEW-1: SESSION GUARD — ป้องกัน duplicate run
# ==========================================
if "is_running" not in st.session_state: st.session_state.is_running = False
if "run_lock" not in st.session_state: st.session_state.run_lock = False

# ==========================================
# 🧮 TOKEN COUNTER STATE
# ==========================================
if "session_token_count" not in st.session_state: st.session_state.session_token_count = 0
if "session_cost_usd" not in st.session_state: st.session_state.session_cost_usd = 0.0
if "total_chars_sent" not in st.session_state: st.session_state.total_chars_sent = 0

# ==========================================
# 💾 BRIEF CACHE — โหลดจากไฟล์
# ==========================================
def load_cache():
    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except: pass
    return {}

def save_cache(c):
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(c, f, ensure_ascii=False, indent=2)

def brief_hash(brief, agents, lang):
    key = f"{brief.strip()}|{sorted(agents)}|{lang}"
    return hashlib.md5(key.encode()).hexdigest()

if "brief_cache" not in st.session_state:
    st.session_state.brief_cache = load_cache()

# ==========================================
# 📊 ANALYTICS — โหลด/บันทึก
# ==========================================
def load_analytics():
    if os.path.exists(ANALYTICS_FILE):
        try:
            with open(ANALYTICS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except: pass
    return {"sessions": [], "agent_usage": {}, "project_count": {}, "weekly": {}}

def save_analytics(a):
    with open(ANALYTICS_FILE, "w", encoding="utf-8") as f:
        json.dump(a, f, ensure_ascii=False, indent=2)

if "analytics" not in st.session_state:
    st.session_state.analytics = load_analytics()

# ==========================================
# 🌏 LANGUAGE TOGGLE STATE
# ==========================================
if "ui_lang" not in st.session_state: st.session_state.ui_lang = "TH"

# ==========================================
# ✏️ CUSTOM AGENT PERSONALITY STATE
# — โหลดจาก agent_personas.json (บันทึกจากหน้า 13) ทุกครั้งที่ start
# ==========================================
PERSONA_FILE = "agent_personas.json"

def load_agent_personas() -> dict:
    if os.path.exists(PERSONA_FILE):
        try:
            with open(PERSONA_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            pass
    return {}

def save_agent_personas(p: dict):
    with open(PERSONA_FILE, "w", encoding="utf-8") as f:
        json.dump(p, f, ensure_ascii=False, indent=2)

if "custom_personas" not in st.session_state:
    # โหลดจากไฟล์ก่อน แล้วค่อย override ด้วย session ถ้ามี
    st.session_state.custom_personas = load_agent_personas()

# ==========================================
# 🔴 RATE LIMIT QUEUE SYSTEM
# ==========================================
_api_semaphore = threading.Semaphore(3)  # max 3 concurrent API calls
_last_call_time = {}
_call_lock = threading.Lock()

def rate_limited_call(fn, *args, **kwargs):
    """Exponential backoff + semaphore สำหรับ Gemini API"""
    max_retries = 3
    base_delay = 2
    with _api_semaphore:
        for attempt in range(max_retries):
            try:
                result = fn(*args, **kwargs)
                return result
            except Exception as e:
                err_str = str(e).lower()
                if "429" in err_str or "quota" in err_str or "rate" in err_str:
                    wait = base_delay * (2 ** attempt)
                    time.sleep(wait)
                    if attempt == max_retries - 1:
                        return f"❌ Rate limit exceeded หลัง {max_retries} retries: {e}"
                else:
                    if attempt == max_retries - 1:
                        return f"❌ Error: {e}"
                    time.sleep(1)
    return "❌ Unknown error"

# ==========================================
# 🧮 TOKEN COUNTER HELPER
# ==========================================
# Gemini 2.5 Flash pricing (approximate): $0.075/1M input tokens, $0.30/1M output tokens
# 1 token ≈ 4 chars (Thai/EN mixed)
def estimate_tokens(text: str) -> int:
    """
    ประมาณ token count สำหรับ mixed Thai/EN text
    - ภาษาอังกฤษ: ~4 chars/token
    - ภาษาไทย: ~2 chars/token (Thai unicode = 1 char แต่มักเป็น 1-2 tokens)
    """
    thai_chars = sum(1 for c in text if '\u0e00' <= c <= '\u0e7f')
    other_chars = len(text) - thai_chars
    return max(1, (thai_chars // 2) + (other_chars // 4))

def update_token_counter(prompt_text: str, response_text: str):
    in_tok = estimate_tokens(prompt_text)
    out_tok = estimate_tokens(response_text)
    st.session_state.session_token_count += in_tok + out_tok
    # Gemini 2.5 Flash: ~$0.075/1M input + $0.30/1M output
    cost = (in_tok / 1_000_000 * 0.075) + (out_tok / 1_000_000 * 0.30)
    st.session_state.session_cost_usd += cost

# --- 🎨 CSS STYLES ---
st.markdown("""
    <style>
    .stApp { background: linear-gradient(135deg, #0b0f19 0%, #1a1a2e 100%); color: #e2e8f0; }
    [data-testid="stSidebar"] { background-color: rgba(15, 23, 42, 0.7); border-right: 1px solid #1e293b; }
    
    div.stButton > button:first-child { background: linear-gradient(90deg, #3b82f6, #8b5cf6); color: white; border: none; border-radius: 8px; box-shadow: 0 4px 15px rgba(59, 130, 246, 0.3); transition: all 0.3s ease; font-weight: bold; }
    div.stButton > button:first-child:hover { transform: translateY(-2px); box-shadow: 0 6px 20px rgba(59, 130, 246, 0.6); }
    
    button[kind="primary"], button[data-testid="baseButton-primary"] {
        padding: 20px !important;
        border-radius: 12px !important;
        background: linear-gradient(90deg, #ff0844 0%, #ffb199 100%) !important;
        box-shadow: 0 10px 30px rgba(255, 8, 68, 0.4) !important;
        border: none !important;
        width: 100% !important;
    }
    button[kind="primary"] div, button[data-testid="baseButton-primary"] div, button[kind="primary"] p, button[data-testid="baseButton-primary"] p {
        font-size: 28px !important;
        font-weight: 900 !important;
        letter-spacing: 1px !important;
        color: white !important;
    }
    button[kind="primary"]:hover, button[data-testid="baseButton-primary"]:hover {
        transform: scale(1.02) !important;
        box-shadow: 0 15px 40px rgba(255, 8, 68, 0.6) !important;
    }
    button[kind="primary"]:disabled, button[data-testid="baseButton-primary"]:disabled {
        background: linear-gradient(90deg, #4b5563, #6b7280) !important;
        box-shadow: none !important; transform: none !important; opacity: 0.6 !important; cursor: not-allowed !important;
    }

    .stTextInput>div>div>input, .stTextArea>div>div>textarea, .stSelectbox>div>div>select { background-color: rgba(30, 41, 59, 0.6) !important; color: #ffffff !important; border-radius: 8px !important; border: 1px solid #334155 !important; }
    .stTextInput>div>div>input:focus, .stTextArea>div>div>textarea:focus { border-color: #3b82f6 !important; box-shadow: 0 0 10px rgba(59, 130, 246, 0.5) !important; }
    
    .centered-title { text-align: center; color: #ffffff; font-weight: bold; padding-bottom: 5px; text-shadow: 0 2px 10px rgba(0,0,0,0.5); font-size: 55px !important; }
    .version-tag { font-size: 0.4em; font-weight: 400; color: #94a3b8; margin-left: 10px; vertical-align: middle; }
    .subtitle { text-align: center; color: #94a3b8; font-size: 14px; margin-bottom: 30px; }
    
    .status-card { background: rgba(30, 41, 59, 0.8); border: 1px solid #334155; padding: 12px; border-radius: 8px; margin-bottom: 10px; font-size: 14px; }
    .status-ok { color: #10b981; font-weight: bold; }
    .status-err { color: #ef4444; font-weight: bold; font-size: 12px;}

    /* ── FIX #1: ปุ่ม Template บรีฟ ── */
    .template-btn button { background: linear-gradient(90deg, #0f766e, #0891b2) !important; color: white !important; font-size: 13px !important; padding: 6px 10px !important; }

    /* ── Rating Stars ── */
    .rating-bar { display: flex; gap: 4px; align-items: center; }
    .star { font-size: 20px; cursor: pointer; color: #334155; }
    .star.filled { color: #f59e0b; }

    /* ── ETA Badge ── */
    .eta-badge { background: rgba(59,130,246,0.15); border: 1px solid #3b82f6; border-radius: 8px; padding: 10px 16px; margin-bottom: 12px; color: #93c5fd; font-size: 14px; }

    /* ── History Card ── */
    .history-card { background: rgba(15,23,42,0.9); border: 1px solid #1e293b; border-radius: 10px; padding: 14px; margin-bottom: 10px; cursor: pointer; transition: border-color 0.2s; }
    .history-card:hover { border-color: #3b82f6; }

    /* ── Debate Round Badge ── */
    .debate-badge { background: linear-gradient(90deg,#7c3aed,#db2777); border-radius: 20px; padding: 4px 14px; font-size: 12px; color: white; display: inline-block; margin-bottom: 8px; }

    .orbit-wrapper { position: fixed; top: 50%; left: 50%; transform: translate(-50%, -50%); width: 600px; height: 600px; z-index: 9999; pointer-events: none; }
    .orbit-ring-inner { position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); width: 340px; height: 340px; border: 2px dashed rgba(59, 130, 246, 0.4); border-radius: 50%; }
    .orbit-ring-outer { position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); width: 500px; height: 500px; border: 2px dashed rgba(139, 92, 246, 0.4); border-radius: 50%; }
    .center-logo { position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); font-size: 80px; text-shadow: 0 0 30px rgba(255,255,255,0.8); z-index: 10000; animation: pulse 2s infinite; }
    .orbit-system { position: absolute; top: 0; left: 0; width: 100%; height: 100%; animation: spin-cw 20s linear infinite; }
    .agent-icon { position: absolute; font-size: 40px; background: rgba(15, 23, 42, 0.95); border-radius: 50%; width: 80px; height: 80px; display: flex; justify-content: center; align-items: center; box-shadow: 0 0 20px rgba(0,0,0,0.8); }
    .agent-outer { border: 3px solid #ffb199; box-shadow: 0 0 25px rgba(255, 177, 153, 0.6); }
    .agent-inner { border: 3px solid #3b82f6; box-shadow: 0 0 25px rgba(59, 130, 246, 0.6); }
    @keyframes spin-cw { from { transform: rotate(0deg); } to { transform: rotate(360deg); } }
    @keyframes spin-ccw { from { transform: translate(-50%, -50%) rotate(0deg); } to { transform: translate(-50%, -50%) rotate(-360deg); } }
    @keyframes pulse { 0% { transform: translate(-50%, -50%) scale(1); } 50% { transform: translate(-50%, -50%) scale(1.1); } 100% { transform: translate(-50%, -50%) scale(1); } }
    .token-card { background: linear-gradient(135deg, rgba(16,185,129,0.1), rgba(5,150,105,0.05)); border: 1px solid rgba(16,185,129,0.3); border-radius: 10px; padding: 12px 16px; margin: 8px 0; }
    .token-num { font-size: 22px; font-weight: 900; color: #10b981; }
    .token-label { font-size: 11px; color: #6ee7b7; letter-spacing: 1px; text-transform: uppercase; }
    .cache-hit { background: linear-gradient(90deg, #f59e0b, #fbbf24); border-radius: 20px; padding: 5px 14px; font-size: 13px; color: #1c1917; font-weight: 800; display: inline-block; margin-bottom: 8px; }
    .kw-chip { display: inline-block; background: rgba(99,102,241,0.2); border: 1px solid rgba(99,102,241,0.4); border-radius: 20px; padding: 3px 10px; margin: 3px; font-size: 12px; color: #a5b4fc; }
    .action-chip { display: inline-block; background: rgba(239,68,68,0.15); border: 1px solid rgba(239,68,68,0.3); border-radius: 20px; padding: 3px 10px; margin: 3px; font-size: 12px; color: #fca5a5; }
    .analytics-card { background: rgba(15,23,42,0.9); border: 1px solid #1e3a5f; border-radius: 12px; padding: 16px; margin-bottom: 12px; }
    .analytics-num { font-size: 28px; font-weight: 900; color: #38bdf8; }
    .analytics-label { font-size: 12px; color: #64748b; }
    .running-banner { background: linear-gradient(90deg, rgba(239,68,68,0.2), rgba(251,146,60,0.2)); border: 1px solid rgba(239,68,68,0.4); border-radius: 10px; padding: 10px 16px; color: #fca5a5; font-weight: bold; text-align: center; margin-bottom: 10px; }
    .retry-badge { background: rgba(245,158,11,0.2); border: 1px solid #f59e0b; border-radius: 8px; padding: 4px 10px; font-size: 12px; color: #fcd34d; display: inline-block; }
    .lang-pill { display: inline-block; background: rgba(99,102,241,0.3); border: 1px solid #6366f1; border-radius: 20px; padding: 3px 12px; font-size: 12px; color: #c7d2fe; margin-left: 6px; }
    .persona-box { background: rgba(30,41,59,0.6); border: 1px dashed #475569; border-radius: 8px; padding: 10px; margin-top: 6px; font-size: 12px; }
    </style>
    """, unsafe_allow_html=True)

st.markdown("""
    <style>
    .marquee-container { background: rgba(255,255,255,0.05); backdrop-filter: blur(10px); padding: 10px; border-radius: 10px; margin-bottom: 20px; border: 1px solid rgba(255,255,255,0.15); box-shadow: 0 4px 15px rgba(0,0,0,0.2); overflow: hidden; white-space: nowrap; width: 100%; position: relative; }
    .marquee-content { display: inline-block; color: white; font-size: 30px; font-weight: 800; letter-spacing: 1.5px; text-shadow: 1px 1px 2px rgba(0,0,0,0.3); animation: marquee-scroll 20s linear infinite; }
    @keyframes marquee-scroll { 0% { transform: translateX(100vw); } 100% { transform: translateX(-100%); } }
    </style>
    <div class="marquee-container">
        <div class="marquee-content">
            AQUALINE QUALITY COMMITTED WITH HEART ❤️ &nbsp;&nbsp;&nbsp;&nbsp; ก็ถ้าไม่ทำอะไร ก็ไม่ต้องพู๊ดดด &nbsp;&nbsp;&nbsp;&nbsp; ❤️ &nbsp;&nbsp;&nbsp;&nbsp; AQUALINE QUALITY COMMITTED WITH HEART ❤️ 
        </div>
    </div>
    """, unsafe_allow_html=True)

st.markdown("""
    <h1 class="centered-title">🎯 AQUALINE STUDIO SPECIAL TEAM <span class="version-tag">V9.0 ULTRA</span></h1>
    <p class="subtitle">Multi-Agent AI · Parallel Processing · Debate Mode · Smart Knowledge · Auto-Retry · Token Counter · PDF Export · AI Keywords</p>
    """, unsafe_allow_html=True)

# ==========================================
# 🗄️ PROJECT VAULT FUNCTIONS
# ==========================================
def load_vault():
    if os.path.exists(VAULT_FILE):
        try:
            with open(VAULT_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
                for k in data:
                    if "knowledge" not in data[k]: data[k]["knowledge"] = ""
                    if "history" not in data[k]: data[k]["history"] = []
                    # ── MEMORY SYSTEM: ต่อยอด vault structure ──
                    if "memory" not in data[k]: data[k]["memory"] = ""
                    if "pinned_facts" not in data[k]: data[k]["pinned_facts"] = []
                    if "memory_updated_at" not in data[k]: data[k]["memory_updated_at"] = ""
                return data
        except: pass
    return {"Default Project": {"url": "", "brief": "", "knowledge": "", "history": [],
                                 "memory": "", "pinned_facts": [], "memory_updated_at": ""}}

def save_vault(vault_data):
    with open(VAULT_FILE, "w", encoding="utf-8") as f:
        json.dump(vault_data, f, ensure_ascii=False, indent=4)

# ==========================================
# 🧠 MEMORY SYSTEM — ฟังก์ชันช่วยจัดการ Memory
# ==========================================
def get_memory_context(project_data: dict) -> str:
    """สร้าง memory context string สำหรับใส่ใน prompt ทุก agent"""
    parts = []
    memory = project_data.get("memory", "").strip()
    pinned = [f.strip() for f in project_data.get("pinned_facts", []) if f.strip()]
    updated_at = project_data.get("memory_updated_at", "")

    if pinned:
        parts.append("📌 ข้อเท็จจริงสำคัญ (Pinned Facts):\n" + "\n".join(f"• {f}" for f in pinned))
    if memory:
        ts_note = f" (อัปเดตล่าสุด: {updated_at})" if updated_at else ""
        parts.append(f"🧠 สรุปความรู้สะสมของ project นี้{ts_note}:\n{memory}")

    if not parts:
        return ""
    return "[PROJECT MEMORY — อ่านก่อนเริ่มทำงาน]\n" + "\n\n".join(parts) + "\n[/PROJECT MEMORY]\n\n"

def save_project_memory(project_name: str, memory_text: str):
    """บันทึก auto-summary memory ลง vault"""
    if project_name not in st.session_state.vault:
        return
    st.session_state.vault[project_name]["memory"] = memory_text.strip()
    st.session_state.vault[project_name]["memory_updated_at"] = datetime.now().strftime("%d/%m/%Y %H:%M")
    save_vault(st.session_state.vault)

def save_pinned_facts(project_name: str, facts: list):
    """บันทึก pinned facts ลง vault"""
    if project_name not in st.session_state.vault:
        return
    cleaned = [f.strip() for f in facts if f.strip()]
    st.session_state.vault[project_name]["pinned_facts"] = cleaned
    save_vault(st.session_state.vault)

def generate_memory_summary(meeting_log: str, brief: str, model_name: str) -> str:
    """เรียก Gemini สรุป meeting log → bullet memory"""
    prompt = f"""อ่านบันทึกการประชุมนี้แล้วสรุปเป็น "ความรู้สะสม" ของ project นี้
ใช้รูปแบบ bullet point กระชับ ไม่เกิน 15 ข้อ
เน้น: ข้อมูลสินค้า, กลุ่มเป้าหมาย, กลยุทธ์ที่ตกลงกัน, ข้อห้าม/ข้อควรระวัง, ผลที่ผ่านมา

บรีฟ: {brief[:300]}

บันทึกประชุม:
{meeting_log[:10000]}

ตอบเป็นภาษาไทย เฉพาะ bullet points ไม่ต้องมีคำนำ"""
    result = "".join(list(call_gemini_true_stream(prompt, model_name, max_output_tokens=2048)))
    return result

if "vault" not in st.session_state: st.session_state.vault = load_vault()
if "current_project" not in st.session_state: st.session_state.current_project = list(st.session_state.vault.keys())[0]


# ==========================================
# 🎮 PIXEL ART CHARACTER GENERATOR
# ==========================================
import base64, os

def get_pixel_art(agent_id: str, size: int = 80) -> str:
    path = os.path.join("agents", f"{agent_id}.png")
    if os.path.exists(path):
        with open(path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        return f'<img src="data:image/png;base64,{b64}" width="{size}" height="{size}" style="image-rendering:pixelated; border-radius:8px"/>'
    # fallback → SVG เดิมถ้าไม่มีรูป
    return f'<span style="font-size:32px">{agent_id}</span>'

# ==========================================
# 👥 TEAM CONFIG — V8: เพิ่ม 4 Agent ใหม่
# ==========================================
@st.cache_resource
def get_team_config():
    return {
        # ── กลุ่มเดิม ──
        "A1": {"name": "นักกลยุทธ์การตลาด",     "icon": "👨‍💼", "p": "วางแผนภาพรวมและจุดขาย",                                          "pixel_art": get_pixel_art("A1"), "parallel": False},
        "A2": {"name": "ผู้จัดการโครงการ",        "icon": "📋",  "p": "คุมเป้าหมายและเวลา",                                              "pixel_art": get_pixel_art("A2"), "parallel": False},
        "A3": {"name": "นักเขียนคำโฆษณา",         "icon": "✍️",  "p": "สร้าง Content และ Caption โซเชียล",                               "pixel_art": get_pixel_art("A3"), "parallel": True},
        "A4": {"name": "กราฟิกดีไซเนอร์",         "icon": "🎨",  "p": "ออกแบบ Visual",                                                   "pixel_art": get_pixel_art("A4"), "parallel": True},
        "A5": {"name": "3D Visualizer",             "icon": "🏗️",  "p": "เรนเดอร์ภาพสินค้าจริง",                                          "pixel_art": get_pixel_art("A5"), "parallel": True},
        "A6": {"name": "ผู้เชี่ยวชาญวิดีโอ",       "icon": "🎬",  "p": "สคริปต์และมุมกล้อง",                                              "pixel_art": get_pixel_art("A6"), "parallel": True},
        "A7": {"name": "นักยิงแอด Facebook",        "icon": "📈",  "p": "วางแผนสื่อโฆษณา",                                                 "pixel_art": get_pixel_art("A7"), "parallel": False},
        "A8": {"name": "ผู้เชี่ยวชาญ SEO",          "icon": "🌐",  "p": "ปรับแต่งเนื้อหาให้ติดอันดับ",                                     "pixel_art": get_pixel_art("A8"), "parallel": True},
        "A9": {"name": "ฝ่ายบริการลูกค้า",         "icon": "💬",  "p": "วางแนวทางตอบคำถาม",                                               "pixel_art": get_pixel_art("A9"), "parallel": True},
        "A10": {"name": "นักวิเคราะห์ข้อมูล",       "icon": "📊",  "p": "วิเคราะห์สถิติความคุ้มค่า",                                       "pixel_art": get_pixel_art("A10"), "parallel": True},
        "A11": {"name": "ครีเอทีฟไดเรกเตอร์",       "icon": "💡",  "p": "คิดไอเดีย Big Idea",                                              "pixel_art": get_pixel_art("A11"), "parallel": False},
        "A12": {"name": "คนเขียนสตอรี่บอร์ด",       "icon": "🎞️",  "p": "วางลำดับภาพเล่าเรื่อง",                                          "pixel_art": get_pixel_art("A12"), "parallel": True},
        "A13": {"name": "อาร์ตไดเรกเตอร์",          "icon": "✨",  "p": "ควบคุมคุณภาพงานดีไซน์",                                          "pixel_art": get_pixel_art("A13"), "parallel": True},
        "A14": {"name": "ผู้เชี่ยวชาญ AI Prompt",   "icon": "🤖",  "p": "ปรับจูนคำสั่งให้ AI",                                             "pixel_art": get_pixel_art("A14"), "parallel": True},
        "A15": {"name": "นักวางระบบอัตโนมัติ",      "icon": "⚙️",  "p": "เชื่อมระบบ Automation",                                           "pixel_art": get_pixel_art("A15"), "parallel": True},
        "A16": {"name": "นักออกแบบบูธ",             "icon": "🎪",  "p": "วางผังงานนิทรรศการ",                                              "pixel_art": get_pixel_art("A16"), "parallel": True},
        "A17": {"name": "นักวิจัยตลาด",             "icon": "🔍",  "p": "เจาะลึกข้อมูลคู่แข่งแบบ Real-time พร้อมอ้างอิง URL",            "pixel_art": get_pixel_art("A17"), "parallel": False},
        "A18": {"name": "ฝ่ายตรวจสเปกสินค้า",       "icon": "✅",  "p": "ตรวจสอบความถูกต้องทางเทคนิค",                                    "pixel_art": get_pixel_art("A18"), "parallel": True},
        "A19": {"name": "นักขายมือโปร",             "icon": "💰",  "p": "สร้างสคริปต์ปิดการขาย",                                          "pixel_art": get_pixel_art("A19"), "parallel": True},
        "A20": {"name": "ที่ปรึกษากฎหมาย",          "icon": "⚖️",  "p": "ตรวจสอบข้อบังคับและลิขสิทธิ์",                                  "pixel_art": get_pixel_art("A20"), "parallel": False},
        "A21": {"name": "นักเขียนบทความและบล็อก",   "icon": "📝",  "p": "เขียนบทความยาว เนื้อหาเชิงลึก และบล็อกสำหรับเว็บไซต์",          "pixel_art": get_pixel_art("A21"), "parallel": True},
        # ── FIX #4:  ใหม่ 4 ตัว ──
        "A22": {"name": "นักวางราคา/Pricing",        "icon": "🧮",  "p": "วิเคราะห์ราคา ตั้ง Promo Bundle และโครงสร้าง Tier ราคา",        "pixel_art": get_pixel_art("A22"), "parallel": True},
        "A23": {"name": "ผู้เชี่ยวชาญ LINE OA/CRM", "icon": "📱",  "p": "วางแผน LINE OA, Broadcast, CRM และ Loyalty Program",           "pixel_art": get_pixel_art("A23"), "parallel": True},
        "A24": {"name": "TikTok & Reels Specialist",  "icon": "🎵",  "p": "สร้าง Hook, Trend, Script สำหรับ TikTok และ Reels โดยเฉพาะ",  "pixel_art": get_pixel_art("A24"), "parallel": True},
        "A25": {"name": "นักจิตวิทยาการตลาด",       "icon": "🧠",  "p": "วิเคราะห์ Psychology ลูกค้า trigger การซื้อ และ Bias ต่างๆ",   "pixel_art": get_pixel_art("A25"), "parallel": False},
    }
team_data = get_team_config()

# ==========================================
# 🔥 FIX #2: Model Priority — ดึงจาก API แล้วเลือกตัวที่ใช้ได้จริง
# ==========================================
@st.cache_data(ttl=300)
def get_best_model(api_key):
    url = f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}"
    try:
        res = requests.get(url, timeout=10)
        if res.status_code == 200:
            available = [
                m['name'] for m in res.json().get('models', [])
                if 'generateContent' in m.get('supportedGenerationMethods', [])
            ]
            # Priority: ใหม่ที่สุดและใช้ได้กับ key นี้จริง
            priorities = [
                'models/gemini-2.5-flash-preview-05-20',
                'models/gemini-2.5-flash-preview-04-17',
                'models/gemini-2.5-flash',
                'models/gemini-2.5-pro-preview-05-06',
                'models/gemini-2.5-pro',
                'models/gemini-2.0-flash',
                'models/gemini-2.0-flash-lite',
                'models/gemini-2.0-flash-001',
                'models/gemini-1.5-flash-latest',
                'models/gemini-1.5-pro-latest',
            ]
            for p in priorities:
                if p in available:
                    return p
            # fallback: เลือกตัวแรกที่มีคำว่า flash หรือ pro
            for m in available:
                if 'flash' in m or 'pro' in m:
                    return m
            return available[0] if available else "models/gemini-2.0-flash"
    except:
        pass
    return "models/gemini-2.0-flash"

# ==========================================
# 🔥 TRUE STREAMING FUNCTION — V9: with retry + rate limit
# ==========================================
def call_gemini_true_stream(prompt, model_name, media_data=None, media_list=None,
                             use_search=False, system_instruction=None,
                             max_output_tokens=None):
    # ── B-1: Dynamic token limit ──────────────────────────────────────────
    # คำนวณ input size แล้วตั้ง output budget ให้เหมาะสม
    # Gemini 2.5 Flash รองรับ output สูงสุด 65,536 tokens
    if max_output_tokens is None:
        input_chars = len(prompt)
        if input_chars < 3000:
            max_output_tokens = 8192    # บรีฟสั้น
        elif input_chars < 8000:
            max_output_tokens = 16384   # บรีฟปานกลาง
        else:
            max_output_tokens = 32768   # บรีฟยาวมาก / context สะสมหนัก
    # ──────────────────────────────────────────────────────────────────────
    url = f"https://generativelanguage.googleapis.com/v1beta/{model_name}:streamGenerateContent?alt=sse&key={API_KEY}"
    parts = [{"text": prompt}]
    if media_list:
        for m in media_list:
            parts.append({"inlineData": {"mimeType": m['type'], "data": m['b64']}})
    elif media_data:
        parts.append({"inlineData": {"mimeType": media_data['type'], "data": media_data['b64']}})
    payload = {
        "contents": [{"parts": parts}],
        "generationConfig": {"temperature": 0.7, "maxOutputTokens": max_output_tokens}
    }
    if system_instruction:
        payload["system_instruction"] = {"parts": [{"text": system_instruction}]}
    if use_search:
        payload["tools"] = [{"googleSearch": {}}]

    max_retries = 3
    for attempt in range(max_retries):
        try:
            # B-2: timeout 240s รองรับ agent ที่ตอบยาว (32K tokens ~120-180s)
            with requests.post(url, json=payload, stream=True, timeout=240) as response:
                if response.status_code == 429:
                    wait_time = 2 * (2 ** attempt)
                    yield f"⏳ Rate limit — รอ {wait_time}s (ครั้งที่ {attempt+1}/{max_retries})..."
                    time.sleep(wait_time)
                    continue
                if response.status_code != 200:
                    if attempt < max_retries - 1:
                        yield f"⚠️ HTTP {response.status_code} — retry {attempt+1}/{max_retries}..."
                        time.sleep(2 * (attempt + 1))
                        continue
                    yield f"❌ Error {response.status_code}: {response.text[:200]}"
                    return
                collected = []
                for line in response.iter_lines():
                    if line:
                        decoded_line = line.decode('utf-8')
                        if decoded_line.startswith("data: "):
                            try:
                                data = json.loads(decoded_line[6:])
                                if 'candidates' in data and len(data['candidates']) > 0:
                                    text_chunk = data['candidates'][0]['content']['parts'][0].get('text', '')
                                    if text_chunk:
                                        collected.append(text_chunk)
                                        yield text_chunk
                            except Exception: pass
                # update token counter
                update_token_counter(prompt, "".join(collected))
                return  # success
        except requests.exceptions.Timeout:
            if attempt < max_retries - 1:
                yield f"⏱️ Timeout — retry {attempt+1}/{max_retries}..."
                time.sleep(3)
            else:
                yield f"❌ Timeout หลัง {max_retries} retry"
        except Exception as e:
            err_str = str(e).lower()
            if ("429" in err_str or "quota" in err_str or "rate" in err_str) and attempt < max_retries - 1:
                wait_time = 2 * (2 ** attempt)
                yield f"⏳ Rate limit — backoff {wait_time}s..."
                time.sleep(wait_time)
            elif attempt < max_retries - 1:
                yield f"⚠️ Error retry {attempt+1}: {str(e)[:80]}..."
                time.sleep(2)
            else:
                yield f"❌ Connection Error หลัง {max_retries} retry: {str(e)[:150]}"

# ─── Non-streaming version สำหรับ parallel — V9: with semaphore ───
def call_gemini_sync(prompt, model_name, media_list=None, use_search=False):
    """Thread-safe sync call with rate limit semaphore"""
    with _api_semaphore:
        result = "".join(list(call_gemini_true_stream(prompt, model_name, media_list=media_list, use_search=use_search)))
        return result

# ==========================================
# FIX #3: SMART KNOWLEDGE CHUNKING
# ==========================================
def smart_chunk_knowledge(text, max_chars=12000):
    """แทนการตัดแบบดิบๆ ด้วย [:8000] → ใช้ paragraph-aware chunking"""
    if len(text) <= max_chars:
        return text
    paragraphs = text.split('\n')
    result, total = [], 0
    for p in paragraphs:
        if total + len(p) > max_chars:
            result.append(f"\n[⚠️ ข้อมูลถูกตัดที่ {max_chars} chars เพื่อประหยัด token — ส่วนที่เหลือไม่ถูกส่ง]")
            break
        result.append(p)
        total += len(p)
    return '\n'.join(result)

# ==========================================
# 🌏 LANGUAGE HELPER
# ==========================================
def lang_suffix(lang="TH"):
    if lang == "TH":
        return "\n\n[สำคัญ: ตอบเป็นภาษาไทยทั้งหมด]"
    else:
        return "\n\n[IMPORTANT: Respond entirely in English.]"

# ==========================================
# 📄 PDF EXPORT — Aqualine Branded
# ==========================================
def generate_branded_pdf(project_name, brief_text, results, summary_text="", ratings=None):
    """สร้าง PDF สวยงามในสไตล์ Aqualine (รองรับภาษาไทยด้วย Sarabun font)"""
    buf = BytesIO()
    if not REPORTLAB_OK:
        return None

    # ── ลงทะเบียน Sarabun font สำหรับภาษาไทย ──
    _FONT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")
    _font_regular = os.path.join(_FONT_DIR, "Sarabun-Regular.ttf")
    _font_bold    = os.path.join(_FONT_DIR, "Sarabun-Bold.ttf")
    _thai_font    = "Helvetica"   # fallback
    _thai_font_b  = "Helvetica-Bold"
    if os.path.exists(_font_regular) and os.path.exists(_font_bold):
        try:
            if "Sarabun" not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont("Sarabun", _font_regular))
            if "Sarabun-Bold" not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont("Sarabun-Bold", _font_bold))
            _thai_font   = "Sarabun"
            _thai_font_b = "Sarabun-Bold"
        except Exception:
            pass  # fallback to Helvetica
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2.5*cm, bottomMargin=2*cm,
        title=f"AQUALINE REPORT — {project_name}"
    )
    BRAND_BLUE   = colors.HexColor("#0a5fa8")
    BRAND_CYAN   = colors.HexColor("#00b4d8")
    BRAND_DARK   = colors.HexColor("#0b0f19")
    BRAND_LIGHT  = colors.HexColor("#e2e8f0")
    ACCENT_RED   = colors.HexColor("#ff0844")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("AqTitle", parent=styles["Normal"],
        fontSize=26, fontName=_thai_font_b,
        textColor=BRAND_BLUE, spaceAfter=4, alignment=TA_CENTER)
    sub_style = ParagraphStyle("AqSub", parent=styles["Normal"],
        fontSize=10, fontName=_thai_font,
        textColor=colors.HexColor("#64748b"), spaceAfter=14, alignment=TA_CENTER)
    h1_style = ParagraphStyle("AqH1", parent=styles["Normal"],
        fontSize=14, fontName=_thai_font_b,
        textColor=BRAND_BLUE, spaceBefore=14, spaceAfter=6,
        borderPad=4, backColor=colors.HexColor("#e0f2fe"),
        leftIndent=6, rightIndent=6)
    h2_style = ParagraphStyle("AqH2", parent=styles["Normal"],
        fontSize=12, fontName=_thai_font_b,
        textColor=colors.HexColor("#0369a1"), spaceBefore=10, spaceAfter=4)
    body_style = ParagraphStyle("AqBody", parent=styles["Normal"],
        fontSize=9.5, fontName=_thai_font,
        textColor=colors.HexColor("#1e293b"), spaceAfter=4, leading=16)
    rating_style = ParagraphStyle("AqRating", parent=styles["Normal"],
        fontSize=9, fontName=_thai_font,
        textColor=colors.HexColor("#475569"), spaceAfter=3)

    story = []

    # ── Header ──
    story.append(Paragraph("🌊 AQUALINE STUDIO", title_style))
    story.append(Paragraph("SPECIAL TEAM — AI MEETING REPORT", sub_style))
    story.append(Paragraph(f"โปรเจกต์: {project_name}  |  วันที่: {datetime.now().strftime('%d/%m/%Y %H:%M')}", sub_style))
    story.append(HRFlowable(width="100%", thickness=2, color=BRAND_BLUE, spaceAfter=10))

    # ── Brief ──
    story.append(Paragraph("📋 บรีฟงาน", h1_style))
    safe_brief = brief_text.replace("\n", "<br/>")
    story.append(Paragraph(safe_brief, body_style))
    story.append(Spacer(1, 8))

    # ── Results per agent ──
    story.append(Paragraph("🤖 ความเห็นจากทีม Special Team", h1_style))
    for agent_name, answer in results:
        story.append(Paragraph(f"▸ {agent_name}", h2_style))
        # BUG5: escape HTML special chars ก่อน ป้องกัน reportlab crash
        safe_ans = answer[:3000] if len(answer) > 3000 else answer
        safe_ans = safe_ans.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        safe_ans = safe_ans.replace("\n", "<br/>")
        if len(answer) > 3000:
            safe_ans += "...<br/><i>[ดูรายละเอียดเพิ่มเติมใน app]</i>"
        try:
            story.append(Paragraph(safe_ans, body_style))
        except Exception:
            story.append(Paragraph("[ไม่สามารถแสดงข้อความนี้ใน PDF]", body_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1"), spaceAfter=6))

    # ── Summary ──
    if summary_text:
        story.append(Paragraph("📌 Action Plan สรุป", h1_style))
        story.append(Paragraph(summary_text.replace("\n", "<br/>"), body_style))

    # ── Ratings ──
    if ratings:
        story.append(Paragraph("⭐ คะแนน Agent", h1_style))
        for rname, rval in ratings.items():
            stars = "★" * rval + "☆" * (5 - rval)
            story.append(Paragraph(f"{rname}: {stars} ({rval}/5)", rating_style))

    # ── Footer ──
    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=1, color=BRAND_CYAN, spaceAfter=6))
    story.append(Paragraph("Generated by AQUALINE STUDIO SPECIAL TEAM V9.0 ULTRA · Powered by Google Gemini AI",
        ParagraphStyle("footer", parent=styles["Normal"],
            fontSize=8, fontName=_thai_font, textColor=colors.HexColor("#94a3b8"), alignment=TA_CENTER)))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()

# ── Budget ──
def load_chairman_budget() -> dict:
    default = {"month": datetime.now().strftime("%Y-%m"), "used_thb": 0.0}
    if os.path.exists(CHAIRMAN_BUDGET_FILE):
        try:
            with open(CHAIRMAN_BUDGET_FILE, "r", encoding="utf-8") as f:
                d = json.load(f)
            if d.get("month") != datetime.now().strftime("%Y-%m"):
                return default
            return d
        except:
            pass
    return default

def save_chairman_budget(used_thb: float):
    with open(CHAIRMAN_BUDGET_FILE, "w", encoding="utf-8") as f:
        json.dump({"month": datetime.now().strftime("%Y-%m"),
                   "used_thb": round(used_thb, 6)}, f)

def get_chairman_budget_used() -> float:
    return load_chairman_budget().get("used_thb", 0.0)

def add_chairman_cost_claude(input_text: str, output_text: str) -> float:
    inp = max(1, len(input_text)  // 4)
    out = max(1, len(output_text) // 4)
    usd = (inp/1_000_000 * CLAUDE_INPUT_USD) + (out/1_000_000 * CLAUDE_OUTPUT_USD)
    thb = round(usd * USD_TO_THB, 6)
    save_chairman_budget(get_chairman_budget_used() + thb)
    return thb

# ── Detect Mode ──
def detect_meeting_mode(brief: str, meeting_log: str) -> str:
    text = (brief + " " + meeting_log[:2000]).lower()
    if any(k in text for k in ["วิเคราะห์แอด","วิเคราะห์ ads","ผลแคมเปญ","ยิงไปแล้ว","ctr","cpm","roas"]):
        return "ads_analyze"
    if any(k in text for k in ["ยิงแอด","ยิงโฆษณา","สร้างแอด","facebook ads","โฆษณา","ad campaign","boost"]):
        return "ads_create"
    if any(k in text for k in ["บทความ","เขียน","content","blog","โพสต์","caption","สคริปต์","script"]):
        return "content"
    if any(k in text for k in ["กลยุทธ์","strategy","แผนการตลาด","branding","positioning"]):
        return "strategy"
    if any(k in text for k in ["สินค้า","product","ราคา","คู่แข่ง","usp","จุดขาย"]):
        return "product"
    return "auto"

# ── Build Prompt ──
def build_chairman_prompt(brief: str, meeting_log: str,
                           custom_instruction: str, lang: str) -> tuple:
    mode = detect_meeting_mode(brief, meeting_log)
    lang_note = "ตอบเป็นภาษาอังกฤษ" if lang == "EN" else "ตอบเป็นภาษาไทย"
    format_map = {
        "ads_create": """โครงสร้างคำตอบ (ห้ามข้าม):
1. 🎯 ประเมินภาพรวม — ดีหรือไม่ดี คะแนน 0-10 พร้อมเหตุผล
2. ✍️ เนื้อหาโพสต์ — Headline / Body copy / CTA พร้อมใช้งานได้เลย
3. 🎪 วัตถุประสงค์โฆษณา — เลือกอะไร (Awareness/Traffic/Lead/Conversion) + เหตุผล
4. 👥 กลุ่มเป้าหมาย — อายุ, interest, behavior, lookalike ที่แนะนำ
5. 📐 รูปแบบโฆษณา — Single/Carousel/Reels/Story + เหตุผล
6. 💰 งบและช่วงเวลา — daily budget, ระยะทดสอบ, schedule
7. 📍 Placement — Feed/Reels/Story/Audience Network + เหตุผล
8. 📊 KPI เป้าหมาย — CTR, CPM, ROAS ที่ควรได้
9. ⚠️ สิ่งที่ต้องแก้ก่อนยิง (ถ้ามี)
10. 🔄 ทางเลือกใหม่ที่ดีกว่า (ถ้าผลออกมาไม่ดีพอ)""",
        "ads_analyze": """โครงสร้างคำตอบ:
1. 📊 สรุปผลแคมเปญ — ตัวเลขที่ได้ vs เป้าหมาย
2. 🔍 Root cause — ทำไมดี/ไม่ดี วิเคราะห์เชิงลึก
3. 🎨 Creative ที่ดีที่สุด — ภาพ/copy ไหนทำงานดีที่สุดและทำไม
4. 👥 Audience insight — กลุ่มไหน response ดีที่สุด
5. ⚡ Action ถัดไป — scale/pause/A-B test อะไร
6. 🗓️ แผน optimization รอบถัดไปแบบละเอียด""",
        "content": """โครงสร้างคำตอบ:
1. ⭐ ประเมินคุณภาพ — clarity, engagement, SEO score
2. 🏗️ โครงสร้างที่ดีกว่า — outline ที่แนะนำ
3. 🪝 Hook/Headline ที่แรงกว่า — 3 ตัวเลือก
4. 📢 CTA — ควรจบด้วยอะไร เพราะอะไร
5. 📱 ช่องทางที่เหมาะ — Facebook/Blog/TikTok/LINE
6. ✏️ ตัวอย่าง draft ที่ปรับปรุงแล้วพร้อมใช้""",
        "strategy": """โครงสร้างคำตอบ:
1. 🗺️ SWOT — จุดแข็ง จุดอ่อน โอกาส ความเสี่ยง
2. 🎯 Positioning — แบรนด์ควรยืนตรงไหนในตลาด
3. ⚡ Priority actions — ทำอะไรก่อน-หลัง เพราะอะไร
4. 📅 Timeline — แผน 30/60/90 วัน
5. 📊 KPI ที่ควรวัดและเป้าหมาย""",
        "product": """โครงสร้างคำตอบ:
1. 💎 USP หลัก — จุดขายที่แข็งแกร่งที่สุด
2. 👤 กลุ่มเป้าหมายที่แท้จริง
3. 💵 ราคา — เหมาะสมหรือควรปรับ เพราะอะไร
4. ⚔️ คู่แข่ง — เราชนะหรือแพ้ตรงไหน
5. 🔧 สิ่งที่ต้องพัฒนา — product/packaging/messaging""",
        "auto": """อ่านบรีฟและบันทึกการประชุมแล้วเลือก format ที่เหมาะที่สุดเอง
ครอบคลุมทุกประเด็นสำคัญ ประเมินว่าดีหรือไม่ดี และเสนอแนวทางใหม่ถ้าผลยังไม่ดีพอ"""
    }
    system_prompt = (
        f"คุณคือ Chairman Master ประธานสูงสุดของ AQUALINE STUDIO SPECIAL TEAM\n"
        f"คุณฉลาดและมีประสบการณ์เหนือกว่าทุก agent ในทีม\n"
        f"หน้าที่: อ่านผลสรุปจากทีมทั้งหมด วิเคราะห์ภาพรวม "
        f"ประเมินว่าดีหรือไม่ดี และเสนอแนวทางที่ดีที่สุด\n"
        f"{lang_note}\n"
        f"{format_map.get(mode, format_map['auto'])}\n"
        f"กฎ: ถ้าผลออกมาไม่ดีพอ ต้องเสนอแนวทางใหม่เสมอ"
    )
    user_prompt = (
        f"{'คำสั่งพิเศษ: ' + custom_instruction + chr(10)*2 if custom_instruction.strip() else ''}"
        f"บรีฟงาน: {brief}\n\n"
        f"บันทึกการประชุมจากทีม:\n{meeting_log[:14000]}"
    )
    return system_prompt, user_prompt

# ── Call Claude Haiku ──
def call_claude_chairman(system_prompt: str, user_prompt: str):
    if not ANTHROPIC_API_KEY:
        yield "❌ ไม่พบ ANTHROPIC_API_KEY\nเพิ่มใน .streamlit/secrets.toml:\nANTHROPIC_API_KEY = 'sk-ant-...'"
        return
    headers = {
        "x-api-key": ANTHROPIC_API_KEY,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    payload = {
        "model": "claude-haiku-4-5",
        "max_tokens": 4096,
        "stream": True,
        "system": system_prompt,
        "messages": [{"role": "user", "content": user_prompt}]
    }
    try:
        with requests.post("https://api.anthropic.com/v1/messages",
                           headers=headers, json=payload,
                           stream=True, timeout=120) as resp:
            if resp.status_code != 200:
                yield f"❌ Claude API Error {resp.status_code}: {resp.text[:300]}"
                return
            for line in resp.iter_lines():
                if not line: continue
                decoded = line.decode("utf-8")
                if decoded.startswith("data: "):
                    try:
                        d = json.loads(decoded[6:])
                        if d.get("type") == "content_block_delta":
                            chunk = d.get("delta", {}).get("text", "")
                            if chunk: yield chunk
                    except: pass
    except requests.exceptions.Timeout:
        yield "\n❌ Timeout — ลองใหม่"
    except Exception as e:
        yield f"\n❌ Error: {str(e)[:200]}"

# ── Call Gemini Chairman ──
def call_gemini_chairman(system_prompt: str, user_prompt: str, model_name: str):
    full_prompt = f"{system_prompt}\n\n{user_prompt}"
    yield from call_gemini_true_stream(full_prompt, model_name, max_output_tokens=8192)

# ==========================================
# 🔍 AI KEYWORD EXTRACTOR
# ==========================================
def extract_keywords_and_actions(meeting_log: str, model_name: str, lang: str = "TH") -> dict:
    """ดึง keyword, action items, deadline จาก meeting log"""
    if lang == "TH":
        extract_prompt = f"""จากบันทึกการประชุมนี้ สกัดออกมาเป็น JSON ดังนี้ (ตอบ JSON อย่างเดียว ไม่ต้องมีข้อความอื่น):
{{
  "keywords": ["keyword1", "keyword2", ...],
  "action_items": ["สิ่งที่ต้องทำ 1", ...],
  "deadlines": ["deadline หรือ timeline ที่พบ", ...],
  "key_decisions": ["การตัดสินใจสำคัญ", ...],
  "risks": ["ความเสี่ยงที่ระบุ", ...]
}}

บันทึกการประชุม:
{meeting_log[:6000]}"""
    else:
        extract_prompt = f"""From this meeting log, extract as JSON only (no other text):
{{
  "keywords": ["keyword1", "keyword2", ...],
  "action_items": ["action 1", ...],
  "deadlines": ["deadlines or timelines found", ...],
  "key_decisions": ["key decision", ...],
  "risks": ["identified risk", ...]
}}

Meeting log:
{meeting_log[:6000]}"""
    raw = "".join(list(call_gemini_true_stream(extract_prompt, model_name)))
    # Strip markdown fences if any
    # BUG6: clean markdown fences อย่างถูกต้อง
    import re as _re
    raw = _re.sub(r"```(?:json)?", "", raw).strip()
    try:
        return json.loads(raw)
    except:
        return {"keywords": [], "action_items": [raw[:500]], "deadlines": [], "key_decisions": [], "risks": []}

# ==========================================
# 📊 ANALYTICS TRACKER
# ==========================================
def record_session_analytics(project: str, agents: list, tokens: int, cost: float, lang: str, brief: str = ""):
    a = st.session_state.analytics
    week_key = datetime.now().strftime("%Y-W%U")
    a["sessions"].append({
        "timestamp": datetime.now().isoformat(),   # ← เพิ่ม field นี้ (หน้า Sessions ใช้)
        "ts": datetime.now().isoformat(),           # ← เก็บเดิมไว้ด้วยเพื่อ backward compat
        "project": project,
        "agents": len(agents),
        "agents_used": agents,                      # ← เพิ่ม list เต็ม (หน้า Sessions/Analytics ใช้)
        "tokens": tokens,
        "cost_usd": round(cost, 5),
        "lang": lang,
        "brief": brief[:120],                       # ← เพิ่ม brief snippet (หน้า Sessions ใช้)
    })
    a["sessions"] = a["sessions"][-200:]  # keep last 200
    for ag in agents:
        a["agent_usage"][ag] = a["agent_usage"].get(ag, 0) + 1
    a["project_count"][project] = a["project_count"].get(project, 0) + 1
    a["weekly"][week_key] = a["weekly"].get(week_key, 0) + 1
    save_analytics(a)

# ==========================================
# FIX #8: BRIEF TEMPLATES
# ==========================================
BRIEF_TEMPLATES = {
    "🎨 ออกแบบแบนเนอร์/กราฟิก": (
        "ขอให้ทีมช่วยวางแผนและออกแบบแบนเนอร์/กราฟิกสำหรับ AQUALINE (https://www.aqualine.co.th/) "
        "โดยระบุ: 1) ขนาดและ Format ที่ต้องการ (Facebook Cover / Story / Feed / GDN Banner) "
        "2) Concept และ Visual Direction ที่สื่อถึงแบรนด์ AQUALINE 3) สี Font และ Element ที่ควรใช้ให้ตรง Brand Guide "
        "4) Copywriting บนชิ้นงาน (Headline / Subhead / CTA) 5) Reference style ที่ควรทำตาม "
        "สิ่งที่ต้องการจากทีม: ไอเดีย Concept, Draft Layout Direction, ข้อความที่ใช้บนชิ้นงาน และ Checklist ก่อนส่งงาน"
    ),
    "📱 คอนเทนต์ Social Media": (
        "ขอให้ทีมช่วยคิดคอนเทนต์โซเชียลมีเดียสำหรับสินค้า/แบรนด์ AQUALINE (https://www.aqualine.co.th/) "
        "โดยครอบคลุม: 1) Hook ที่ดึงดูดใน 3 วินาทีแรก 2) โครงสร้างโพสต์ (Caption + Visual Direction) "
        "3) เนื้อหาที่เหมาะกับแต่ละแพลตฟอร์ม: Facebook, Instagram, TikTok, LINE OA "
        "4) Hashtag ที่เหมาะสม 5) CTA ที่ชัดเจน 6) ไทม์ไลน์การโพสต์ที่แนะนำ "
        "สิ่งที่ต้องการจากทีม: ไอเดียคอนเทนต์อย่างน้อย 5 ชิ้น พร้อม Caption สำเร็จรูปใช้ได้เลย"
    ),
    "🎯 แคมเปญยิง Ads": (
        "ขอให้ทีมช่วยวางแผนแคมเปญโฆษณา Facebook/Instagram Ads สำหรับ AQUALINE (https://www.aqualine.co.th/) "
        "โดยวิเคราะห์และวางแผน: 1) กลุ่มเป้าหมาย (Audience) พร้อม Interest, Behavior, Demographic ที่แนะนำ "
        "2) โครงสร้าง Campaign / Ad Set / Ad ที่เหมาะสม 3) Hook และ Copy สำหรับ Ad Text, Headline, Description "
        "4) รูปแบบโฆษณาที่แนะนำ (Static / Carousel / Video / Story) 5) Budget และ Bidding Strategy "
        "6) KPI ที่ควรวัด (CTR, CPM, ROAS, CPA) และเป้าหมายที่ควรตั้ง "
        "สิ่งที่ต้องการ: แผนแคมเปญพร้อมใช้งาน, Ad Copy สำเร็จรูป, และ Targeting Checklist"
    ),
    "🔍 วิจัยคีย์เวิร์ด + SEO": (
        "ขอให้ทีมช่วยวิเคราะห์และวางแผน SEO / Keyword Strategy สำหรับเว็บไซต์ AQUALINE (https://www.aqualine.co.th/) "
        "โดยครอบคลุม: 1) คีย์เวิร์ดหลักและ Long-tail ที่ควรโฟกัส พร้อมอธิบายว่าทำไม "
        "2) วิเคราะห์คู่แข่งในธุรกิจเดียวกัน: พวกเขาใช้คีย์เวิร์ดอะไร, จุดแข็ง/อ่อนของพวกเขา "
        "3) โครงสร้างบทความ/หน้าเว็บที่ควรสร้างเพื่อติดอันดับ 4) On-page SEO Checklist สำหรับแต่ละหน้า "
        "5) หัวข้อบล็อกที่แนะนำ 10 หัวข้อพร้อม Outline "
        "สิ่งที่ต้องการ: Keyword List, Competitor Gap Analysis, Content Plan สำหรับ 3 เดือน"
    ),
    "✍️ เขียนบล็อก / บทความ": (
        "ขอให้ทีมช่วยวางแผนและเขียนบทความ/บล็อกสำหรับเว็บไซต์ AQUALINE (https://www.aqualine.co.th/) "
        "โดยบทความต้องมี: 1) หัวข้อที่ติด Google และตอบโจทย์ลูกค้ากลุ่มเป้าหมาย "
        "2) โครงสร้างบทความที่ครบ: Intro Hook, Body (แบ่ง Section ชัด), Conclusion + CTA "
        "3) SEO On-page: Meta Title, Meta Description, Alt Text สำหรับรูป, Internal Link Suggestions "
        "4) ความยาวที่เหมาะสม (แนะนำ 1,000-2,000 คำ) 5) Tone of Voice ที่สอดคล้องกับแบรนด์ AQUALINE "
        "สิ่งที่ต้องการ: บทความฉบับร่างพร้อม SEO Tags และคำแนะนำในการปรับก่อนเผยแพร่"
    ),
    "🎬 สคริปต์วิดีโอ + Storyboard": (
        "ขอให้ทีมช่วยเขียนสคริปต์และวางแผน Storyboard สำหรับวิดีโอ/Reels ของ AQUALINE (https://www.aqualine.co.th/) "
        "โดยครอบคลุม: 1) Hook 3 วินาทีแรกที่หยุดนิ้วได้ 2) โครงสร้างวิดีโอ: Scene-by-scene พร้อม Duration "
        "3) สคริปต์พูด (ถ้ามี Presenter) หรือ Text Overlay บนจอ 4) มุมกล้อง, แสง, และ Prop ที่ต้องเตรียม "
        "5) เพลง/Sound Effect ที่เหมาะกับอารมณ์วิดีโอ 6) CTA ตอนท้ายวิดีโอ "
        "สิ่งที่ต้องการ: สคริปต์พร้อมถ่าย, Shot List, และ Checklist วันถ่ายทำ"
    ),
    "🏗️ Mockup / 3D Sketch-up": (
        "ขอให้ทีมช่วยวางแผนและ Briefing สำหรับการทำ 3D Mockup / Sketch-up ของสินค้าหรือพื้นที่ AQUALINE "
        "(https://www.aqualine.co.th/) โดยระบุ: 1) วัตถุประสงค์ของ Mockup: ใช้สำหรับ Presentation, Ads, หรือ Proposal "
        "2) มุมมอง (Camera Angle) ที่ต้องการ: Front, Perspective, Hero Shot "
        "3) Material, สี, และ Texture ที่ต้องการให้แสดง 4) Environment / Background ของ Scene "
        "5) ไฟล์ที่ต้องส่งมอบ: Format, ขนาด, และ Resolution "
        "สิ่งที่ต้องการ: Creative Brief พร้อมส่งให้ทีม 3D, Reference Board, และ Revision Checklist"
    ),
    "🖨️ สิ่งพิมพ์ / Signage": (
        "ขอให้ทีมช่วยวางแผนและออกแบบสิ่งพิมพ์หรือป้ายสำหรับ AQUALINE (https://www.aqualine.co.th/) "
        "เช่น Leaflet, Brochure, Roll-up Banner, Poster, Catalogue, หรือป้ายในงาน Event "
        "โดยครอบคลุม: 1) ขนาดและ Format มาตรฐานสำหรับโรงพิมพ์ (Bleed, Safe Zone, Resolution 300 DPI) "
        "2) Layout Direction และลำดับ Hierarchy ของข้อมูล 3) เนื้อหาที่ต้องอยู่บนชิ้นงาน (ข้อความ, ภาพ, โลโก้) "
        "4) สีที่ใช้ตาม Brand Guide (CMYK สำหรับพิมพ์) 5) ข้อควรระวังก่อนส่งโรงพิมพ์ "
        "สิ่งที่ต้องการ: Layout Draft Direction, Copywriting บนชิ้นงาน และ Print-ready Checklist"
    ),
    "🌐 อัปเดต/ดูแลเว็บไซต์": (
        "ขอให้ทีมช่วยวางแผนการอัปเดตหรือปรับปรุงเว็บไซต์ AQUALINE (https://www.aqualine.co.th/) "
        "โดยวิเคราะห์และแนะนำ: 1) หน้าเว็บหรือ Section ที่ควรอัปเดตเนื้อหา 2) UX/UI ที่ควรปรับให้ดีขึ้น "
        "3) เนื้อหาใหม่ที่ควรเพิ่ม (Landing Page, Product Page, Blog) 4) Speed Optimization และ Technical SEO "
        "5) CTA Placement ที่จะช่วยเพิ่ม Conversion 6) ภาพและวิดีโอที่ควรถ่ายเพิ่มสำหรับหน้าเว็บ "
        "สิ่งที่ต้องการ: To-do List การอัปเดต, เนื้อหาร่างพร้อมใช้, และ Priority Matrix"
    ),
}

# ==========================================
# 👈 SIDEBAR
# ==========================================
with st.sidebar:
    st.header("🗄️ Project Vault")
    selected_project = st.selectbox("📂 เลือกแฟ้มงาน:", list(st.session_state.vault.keys()), index=list(st.session_state.vault.keys()).index(st.session_state.current_project))
    if selected_project != st.session_state.current_project:
        st.session_state.current_project = selected_project; st.rerun()
        
    new_project_name = st.text_input("➕ สร้างแฟ้มใหม่:", placeholder="เช่น Overview-180")
    if st.button("💾 บันทึกแฟ้มใหม่", use_container_width=True):
        if new_project_name and new_project_name not in st.session_state.vault:
            st.session_state.vault[new_project_name] = {"url": "", "brief": "", "knowledge": "", "history": []}
            st.session_state.current_project = new_project_name
            save_vault(st.session_state.vault); st.success("สร้างแฟ้มสำเร็จ!"); time.sleep(1); st.rerun()
            
    st.markdown("---")
    st.header("👥 Multi-Agent Team")
    
    # FIX: Toggle all
    def toggle_all():
        for aid in team_data.keys(): st.session_state[aid] = st.session_state.all_btn
    st.checkbox("✅ เลือกทีมงานทั้งหมด", key="all_btn", on_change=toggle_all)
    
    # ── CSS สำหรับ pixel art agent cards ──
    st.markdown("""
    <style>
    .px-agent-grid { display: flex; flex-wrap: wrap; gap: 6px; margin-bottom: 8px; }
    .px-agent-card {
        display: flex; flex-direction: column; align-items: center;
        background: rgba(15,23,42,0.8);
        border: 1px solid rgba(59,130,246,0.25);
        border-radius: 8px; padding: 6px 4px 4px;
        width: 100px; position: relative;
    }
    .px-agent-card.selected { border-color: #3b82f6; background: rgba(30,58,138,0.5); box-shadow: 0 0 10px rgba(59,130,246,0.4); }
    .px-agent-card.new-agent { border-color: rgba(167,139,250,0.5); }
    .px-agent-card.new-agent.selected { border-color: #a78bfa; box-shadow: 0 0 10px rgba(167,139,250,0.4); }
    .px-agent-name { font-size: 15px; color: #94a3b8; text-align: center; margin-top: 3px; line-height: 1.2; max-width: 100%; word-break: break-word; }
    .px-agent-card.selected .px-agent-name { color: #93c5fd; }
    .px-new-badge { position: absolute; top: -4px; right: -4px; background: #7c3aed; color: white; font-size: 7px; padding: 1px 4px; border-radius: 4px; font-weight: bold; }
    .px-section-label { font-size: 11px; color: #64748b; margin: 8px 0 4px; letter-spacing: 0.5px; }
    .px-agent-card svg { width: 100% !important; height: auto !important; max-height: 90px; }
    .px-agent-card img { width: 100% !important; height: auto !important; max-height: 90px; }
    </style>
                
    """, unsafe_allow_html=True)

    # กลุ่มหลัก — pixel art grid
    st.markdown('<div class="px-section-label">── กลุ่มหลัก ──</div>', unsafe_allow_html=True)
    main_agents = ["A1","A2","A3","A4","A5","A6","A7","A8","A9","A10","A11","A12","A13","A14","A15","A16","A17","A18","A19","A20","A21"]
    new_agents  = ["A22","A23","A24","A25"]

    cols_main = st.columns(3)
    for idx, aid in enumerate(main_agents):
        info = team_data[aid]
        svg  = info.get('pixel_art', '')
        with cols_main[idx % 3]:
            is_checked = st.session_state.get(aid, False)
            sel_cls = "selected" if is_checked else ""
            st.markdown(f"""<div class="px-agent-card {sel_cls}">{svg}<div class="px-agent-name">{info['name']}</div></div>""", unsafe_allow_html=True)
            st.checkbox("", key=aid, label_visibility="collapsed")

    st.markdown('<div class="px-section-label">── ✨ ใหม่ใน V8.0 ──</div>', unsafe_allow_html=True)
    cols_new = st.columns(3)
    for idx, aid in enumerate(new_agents):
        info = team_data[aid]
        svg  = info.get('pixel_art', '')
        with cols_new[idx % 3]:
            is_checked = st.session_state.get(aid, False)
            sel_cls = "selected" if is_checked else ""
            st.markdown(f"""<div class="px-agent-card new-agent {sel_cls}"><span class="px-new-badge">NEW</span>{svg}<div class="px-agent-name">{info['name']}</div></div>""", unsafe_allow_html=True)
            st.checkbox("", key=aid, label_visibility="collapsed")

    selected_agents = [aid for aid in team_data.keys() if st.session_state.get(aid, False)]

    st.markdown("---")
    # FIX #6: Parallel Mode toggle
    st.header("⚡ โหมดการประชุม")
    parallel_mode = st.toggle("🚀 Parallel Mode (เร็ว 3x)", value=False, help="Agent ที่ไม่ depend กันจะทำงานพร้อมกัน ประหยัดเวลามาก แต่ไม่เห็น live stream")
    debate_mode = st.toggle("🥊 Debate Mode (Round 2)", value=False, help="หลังประชุมรอบแรก ให้ Agent comment งานกันและกัน")

    # ── 🌏 V9: Language Toggle ──
    st.markdown("---")
    st.header("🌏 ภาษาของ AI")
    lang_choice = st.radio("AI ตอบเป็นภาษา:", ["🇹🇭 ไทย (TH)", "🇬🇧 English (EN)"],
        index=0 if st.session_state.ui_lang == "TH" else 1,
        help="เลือกภาษาที่ต้องการให้ AI ตอบ — ใช้ EN สำหรับลูกค้าต่างชาติ")
    st.session_state.ui_lang = "TH" if "TH" in lang_choice else "EN"
    st.markdown(f'<span class="lang-pill">{"🇹🇭 ภาษาไทย" if st.session_state.ui_lang == "TH" else "🇬🇧 English"}</span>', unsafe_allow_html=True)

    # ── ✏️ V9: Custom Agent Personality ──
    st.markdown("---")
    st.header("✏️ Custom Agent Persona")
    with st.expander("🎭 ปรับ Personality ของ Agent", expanded=False):
        st.caption("แก้ system prompt สำหรับแต่ละ agent ได้ที่นี่ (เว้นว่างเพื่อใช้ default)")
        if selected_agents:
            for aid in selected_agents:  # B-6: แสดงทุก agent ที่เลือก
                info = team_data[aid]
                default_p = info['p']
                custom_key = f"persona_{aid}"
                current_val = st.session_state.custom_personas.get(aid, "")
                new_val = st.text_area(
                    f"{info['icon']} {info['name']}",
                    value=current_val,
                    placeholder=f"Default: {default_p}",
                    height=60,
                    key=f"persona_input_{aid}"
                )
                if new_val != current_val:
                    st.session_state.custom_personas[aid] = new_val
                    save_agent_personas(st.session_state.custom_personas)
        else:
            st.caption("เลือก Agent ก่อนเพื่อตั้งค่า Persona")

# ==========================================
# 🏠 LAYOUT
# ==========================================
col_main, col_health = st.columns([3.5, 1.2])

if "health_report" not in st.session_state: st.session_state.health_report = {}
if "meeting_log" not in st.session_state: st.session_state.meeting_log = ""
if "raw_results" not in st.session_state: st.session_state.raw_results = []
if "summary_text" not in st.session_state: st.session_state.summary_text = ""
if "ready_to_export" not in st.session_state: st.session_state.ready_to_export = False
if "agent_ratings" not in st.session_state: st.session_state.agent_ratings = {}

if "chairman_result"    not in st.session_state: st.session_state.chairman_result    = ""
if "chairman_cost_thb"  not in st.session_state: st.session_state.chairman_cost_thb  = 0.0
if "chairman_ai_mode"   not in st.session_state: st.session_state.chairman_ai_mode   = "gemini"

with col_main:
    st.subheader(f"📝 รายละเอียดงาน: {st.session_state.current_project}")
    current_data = st.session_state.vault[st.session_state.current_project]

    target_link = st.text_input("🔗 ลิงก์อ้างอิงสินค้า (URL):", value=current_data.get("url", ""))

    # ==========================================
    # 🧠 MEMORY PANEL — แสดงและจัดการ Project Memory
    # ==========================================
    _mem = current_data.get("memory", "").strip()
    _pinned = current_data.get("pinned_facts", [])
    _mem_updated = current_data.get("memory_updated_at", "")

    # badge แจ้งเตือนถ้า memory เพิ่งอัปเดต
    if st.session_state.get("memory_just_updated"):
        st.success("🧠 Memory อัปเดตแล้ว! Agent ทุกคนจะจำข้อมูลนี้ใน session หน้า")
        del st.session_state["memory_just_updated"]

    with st.expander(
        f"🧠 Project Memory {'✅ มีข้อมูลสะสม' if (_mem or _pinned) else '(ยังว่างเปล่า)'}",
        expanded=bool(_mem or _pinned)
    ):
        mem_tab1, mem_tab2 = st.tabs(["📌 Pinned Facts", "🗂️ Auto-Summary Memory"])

        # ── Tab 1: Pinned Facts ──
        with mem_tab1:
            st.caption("ปักหมุดข้อเท็จจริงสำคัญ — agent ทุกคนจะรู้ข้อมูลนี้ทันทีตั้งแต่เริ่ม session")
            # แสดง facts ที่มีอยู่
            current_facts = _pinned if _pinned else [""]
            facts_text = st.text_area(
                "ใส่ข้อมูลสำคัญ (1 บรรทัด = 1 ข้อ)",
                value="\n".join(current_facts),
                height=150,
                placeholder="เช่น:\nสินค้า: น้ำกรอง Aqualine ราคา 890 บาท\nกลุ่มเป้าหมาย: แม่บ้านอายุ 30-45\nข้อห้าม: ห้ามเปรียบเทียบคู่แข่งโดยตรง\nโปร: ซื้อ 2 ลด 15%",
                key="pinned_facts_input"
            )
            col_pin1, col_pin2 = st.columns(2)
            with col_pin1:
                if st.button("💾 บันทึก Pinned Facts", use_container_width=True):
                    new_facts = [f.strip() for f in facts_text.split("\n") if f.strip()]
                    save_pinned_facts(st.session_state.current_project, new_facts)
                    st.success(f"✅ บันทึก {len(new_facts)} ข้อเรียบร้อย!")
                    st.rerun()
            with col_pin2:
                if st.button("🗑️ ล้าง Pinned Facts", use_container_width=True):
                    save_pinned_facts(st.session_state.current_project, [])
                    st.success("ล้างแล้ว")
                    st.rerun()

        # ── Tab 2: Auto-Summary Memory ──
        with mem_tab2:
            if _mem:
                st.caption(f"🕐 อัปเดตล่าสุด: {_mem_updated}" if _mem_updated else "")
                st.markdown(
                    f"<div style='background:rgba(30,41,59,0.6);border:1px dashed #334155;"
                    f"border-radius:8px;padding:12px;font-size:13px;line-height:1.7'>"
                    f"{_mem.replace(chr(10), '<br>')}</div>",
                    unsafe_allow_html=True
                )
                col_m1, col_m2 = st.columns(2)
                with col_m1:
                    if st.button("🔄 สร้าง Memory ใหม่จาก Session ล่าสุด", use_container_width=True):
                        _last_log = ""
                        _hist = current_data.get("history", [])
                        if _hist:
                            _last_log = _hist[-1].get("log", "")
                        if _last_log:
                            with st.spinner("🧠 กำลังสรุป..."):
                                _new_mem = generate_memory_summary(
                                    _last_log,
                                    current_data.get("brief", ""),
                                    get_best_model(API_KEY)
                                )
                                if _new_mem and not _new_mem.startswith("❌"):
                                    save_project_memory(st.session_state.current_project, _new_mem)
                                    st.success("✅ Memory อัปเดตแล้ว!")
                                    st.rerun()
                        else:
                            st.warning("ยังไม่มีประวัติการประชุมสำหรับ project นี้")
                with col_m2:
                    if st.button("🗑️ ล้าง Memory", use_container_width=True):
                        save_project_memory(st.session_state.current_project, "")
                        st.success("ล้างแล้ว")
                        st.rerun()
                # แก้ไข memory ด้วยตนเอง
                with st.expander("✏️ แก้ไข Memory ด้วยตนเอง"):
                    edited_mem = st.text_area("แก้ไขเนื้อหา Memory:", value=_mem, height=200, key="edit_mem_ta")
                    if st.button("💾 บันทึกการแก้ไข", use_container_width=True):
                        save_project_memory(st.session_state.current_project, edited_mem)
                        st.success("✅ บันทึกแล้ว!")
                        st.rerun()
            else:
                st.info("🧠 ยังไม่มี Memory — รันประชุมครั้งแรกแล้วระบบจะสร้าง Memory อัตโนมัติ\nหรือกดปุ่มด้านล่างถ้ามี session เก่าอยู่แล้ว")
                if st.button("🧠 สร้าง Memory จาก Session ล่าสุด", use_container_width=True):
                    _hist = current_data.get("history", [])
                    if _hist:
                        _last_log = _hist[-1].get("log", "")
                        with st.spinner("🧠 กำลังสรุป..."):
                            _new_mem = generate_memory_summary(
                                _last_log,
                                current_data.get("brief", ""),
                                get_best_model(API_KEY)
                            )
                            if _new_mem and not _new_mem.startswith("❌"):
                                save_project_memory(st.session_state.current_project, _new_mem)
                                st.success("✅ สร้าง Memory เรียบร้อย!")
                                st.rerun()
                    else:
                        st.warning("ยังไม่มีประวัติการประชุมสำหรับ project นี้")
    st.markdown("**📋 Template บรีฟด่วน:**")
    tmpl_cols = st.columns(3)
    for i, (tmpl_name, tmpl_text) in enumerate(BRIEF_TEMPLATES.items()):
        with tmpl_cols[i % 3]:
            if st.button(tmpl_name, use_container_width=True, key=f"tmpl_{i}"):
                st.session_state["brief_prefill"] = tmpl_text
                st.rerun()

    prompt_input = st.text_area(
        "💬 พิมพ์สเปกหรือคำสั่งที่นี่:",
        value=st.session_state.get("brief_prefill", current_data.get("brief", "")),
        height=180
    )
    if "brief_prefill" in st.session_state:
        del st.session_state["brief_prefill"]

    # ── AI Knowledge System ──
    st.markdown("##### 👁️ AI Knowledge System (รองรับรูปภาพหลายใบ, PDF, Word)")
    saved_knowledge = current_data.get("knowledge", "")
    if saved_knowledge: st.info(f"📚 มีเอกสารอ้างอิงบันทึกไว้แล้ว ({len(saved_knowledge):,} chars)")

    # FIX #1: รับหลายรูป
    uploaded_files = st.file_uploader("แนบไฟล์ข้อมูล (อัปโหลดได้หลายไฟล์พร้อมกัน)", type=["jpg", "jpeg", "png", "pdf", "docx"], accept_multiple_files=True)
    image_list = []  # FIX #1: เปลี่ยนจาก single image_dict → list
    doc_text = ""
    
    if uploaded_files:
        for file in uploaded_files:
            if file.type in ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                try:
                    if file.type == "application/pdf":
                        with pdfplumber.open(file) as pdf:
                            doc_text += f"\n--- {file.name} ---\n" + "\n".join([p.extract_text() or "" for p in pdf.pages])
                    else:
                        doc = Document(file)
                        doc_text += f"\n--- {file.name} ---\n" + "\n".join([p.text for p in doc.paragraphs])
                    st.success(f"✅ โหลดข้อมูล {file.name} สำเร็จ")
                except Exception as e:
                    st.error(f"❌ Error {file.name}: {str(e)}")
            else:
                # FIX #1: เก็บทุกรูปใน list
                image_list.append({"type": file.type, "b64": base64.b64encode(file.read()).decode('utf-8')})
                st.success(f"✅ โหลดภาพ {file.name} สำเร็จ")
        if image_list:
            st.info(f"🖼️ โหลดรูปภาพทั้งหมด {len(image_list)} ใบ จะส่งให้ทุก Agent ที่เลือก")

    st.markdown("<br>", unsafe_allow_html=True)

    # FIX #5: Pre-flight check + ETA
    if selected_agents:
        est_seconds = len(selected_agents) * (8 if not parallel_mode else 3)
        est_str = f"{est_seconds // 60} นาที {est_seconds % 60} วินาที" if est_seconds >= 60 else f"~{est_seconds} วินาที"
        mode_label = "⚡ Parallel" if parallel_mode else "🔄 Sequential"
        lang_label = f"🌏 {st.session_state.ui_lang}"
        st.markdown(f"""<div class="eta-badge">
            🤖 เลือก {len(selected_agents)} Agent · {mode_label} · {lang_label} · ⏱️ คาดว่าใช้เวลา {est_str}
        </div>""", unsafe_allow_html=True)

    # 🔴 V9: Session Guard — แสดง banner ถ้า running
    if st.session_state.is_running:
        st.markdown('<div class="running-banner">🔴 ระบบกำลังทำงาน... กรุณารอจนเสร็จก่อนกด Run อีกครั้ง</div>', unsafe_allow_html=True)

    # 💾 V9: Cache Check — เช็คก่อนรันว่า brief นี้เคยรันแล้วหรือยัง
    # BUG3: rename เป็น session_cache_key เพื่อไม่ชนกับ variable อื่น
    session_cache_key = brief_hash(prompt_input.strip() if prompt_input else "", selected_agents, st.session_state.ui_lang)
    cached_result = st.session_state.brief_cache.get(session_cache_key)
    if cached_result and not st.session_state.is_running:
        st.markdown('<span class="cache-hit">⚡ พบ Cache — บรีฟนี้เคยรันแล้ว!</span>', unsafe_allow_html=True)
        col_cache1, col_cache2 = st.columns(2)
        with col_cache1:
            if st.button("📂 โหลดผลเดิมทันที (ประหยัด API)", use_container_width=True):
                st.session_state.meeting_log = cached_result["log"]
                st.session_state.raw_results = [tuple(r) for r in cached_result["results"]]
                st.session_state.summary_text = cached_result.get("summary", "")
                st.session_state.health_report = {r[0]: "✅ (Cached)" for r in cached_result["results"]}
                st.rerun()
        with col_cache2:
            if st.button("🔄 รันใหม่ (ข้ามแคช)", use_container_width=True):
                del st.session_state.brief_cache[session_cache_key]
                st.rerun()

    # ── 🔴 RUN TEAM — V9 with Session Guard ──
    run_disabled = st.session_state.is_running
    if st.button("🚀 เริ่มระดมสมอง SPECIAL TEAM", type="primary", use_container_width=True, disabled=run_disabled):
        if not selected_agents:
            st.error("⚠️ กรุณาเลือกทีมงานก่อนครับ")
        elif not prompt_input.strip():
            st.error("⚠️ กรุณาพิมพ์บรีฟงานหรือเลือก Template ก่อนครับ")
        else:
            # 🔐 ล็อก session ทันที
            st.session_state.is_running = True
            # รีเซ็ต token counter สำหรับ session ใหม่
            st.session_state.session_token_count = 0
            st.session_state.session_cost_usd = 0.0
            _run_error = None  # BUG1: track error เพื่อ finally unlock
            # บันทึก Vault
            st.session_state.vault[st.session_state.current_project]["url"] = target_link
            st.session_state.vault[st.session_state.current_project]["brief"] = prompt_input
            if doc_text:
                st.session_state.vault[st.session_state.current_project]["knowledge"] = doc_text
            save_vault(st.session_state.vault)
            
            st.session_state.meeting_log = ""
            st.session_state.raw_results = []
            st.session_state.summary_text = ""
            st.session_state.health_report = {}
            st.session_state.agent_ratings = {}

            # ── Orbit Animation (Pixel Art) ──
            num = len(selected_agents)
            icons_html = ""
            radius_inner, radius_outer = 170, 250
            center_x, center_y = 300, 300
            for i, aid in enumerate(selected_agents):
                angle = (2 * math.pi / num) * i
                r = radius_outer if i % 2 == 0 else radius_inner
                style_class = "agent-outer" if i % 2 == 0 else "agent-inner"
                x = center_x + r * math.cos(angle)
                y = center_y + r * math.sin(angle)
                px_svg = team_data[aid].get('pixel_art', team_data[aid]['icon'])
                icons_html += f'<div class="agent-icon {style_class}" style="left:{x}px; top:{y}px; animation: spin-ccw 20s linear infinite; font-size:14px;">{px_svg}</div>'
            orbit_html = f"""
            <div class="orbit-wrapper">
            <div class="orbit-ring-inner"></div>
            <div class="orbit-ring-outer"></div>
            <div class="center-logo">
             <img src="https://static.wixstatic.com/media/1af48d_3d1d8bf08172488c950dbceb3c0dbe2f~mv2.png" 
             style="width: 120px; height: 120px; border-radius: 300%; object-fit: contain; background: white;">
            </div>
            <div class="orbit-system">{icons_html}</div>
            </div>"""
            orbit_place = st.empty()
            orbit_place.markdown(orbit_html, unsafe_allow_html=True)

            active_model = get_best_model(API_KEY)
            st.info(f"🔮 ใช้โมเดล: **{active_model}** · 🌏 ภาษา: **{st.session_state.ui_lang}**")
            status = st.status(f"🔮 ระบบกำลังวิเคราะห์...")

            # FIX #3: Smart chunking
            raw_knowledge = doc_text if doc_text else saved_knowledge
            final_knowledge = smart_chunk_knowledge(raw_knowledge, max_chars=12000)
            meeting_context = ""
            stream_container = st.empty()
            _lang = st.session_state.ui_lang
            _lang_sfx = lang_suffix(_lang)

            def build_prompt(aid, name, icon, p, meeting_ctx, use_search=False):
                """สร้าง prompt โดยรองรับ custom persona + language + PROJECT MEMORY"""
                custom_p = st.session_state.custom_personas.get(aid, "").strip()
                role_desc = custom_p if custom_p else p
                search_note = " (สำคัญมาก: ค้นหาข้อมูลที่เป็นปัจจุบัน และพิมพ์ลิงก์ URL แหล่งอ้างอิงด้วย)" if use_search else ""
                # BUG2: จำกัด meeting_ctx ที่ส่งใน prompt ไม่ให้เกิน 8000 chars
                safe_ctx = meeting_ctx[-8000:] if len(meeting_ctx) > 8000 else meeting_ctx
                # ── MEMORY: ดึง context จาก vault ใส่ขึ้นต้น prompt ──
                mem_ctx = get_memory_context(st.session_state.vault[st.session_state.current_project])
                return (f"{mem_ctx}"
                        f"คุณคือ {name} ({role_desc}){search_note}\n"
                        f"แฟ้มงาน: {st.session_state.current_project}\n"
                        f"บรีฟ: {prompt_input}\nลิงก์: {target_link}\n"
                        f"[ข้อมูล]:\n{final_knowledge}\n"
                        f"ความเห็นก่อนหน้า: {safe_ctx}"
                        f"{_lang_sfx}")

            # ─────────────────────────────────────────
            # FIX #6: PARALLEL vs SEQUENTIAL mode
            # ─────────────────────────────────────────
            try:  # BUG1: try/finally ป้องกัน is_running ค้างถ้า error
                if parallel_mode:
                    # แบ่ง agent เป็น parallel-safe และ sequential
                    parallel_ids = [aid for aid in selected_agents if team_data[aid].get("parallel", True)]
                    sequential_ids = [aid for aid in selected_agents if not team_data[aid].get("parallel", True)]

                    # รัน sequential ก่อน
                    for aid in sequential_ids:
                        name, icon, p = team_data[aid]['name'], team_data[aid]['icon'], team_data[aid]['p']
                        use_search = (aid == "A17")
                        multi_prompt = build_prompt(aid, name, icon, p, meeting_context, use_search)
                        status.update(label=f"{icon} {name} กำลังวิเคราะห์... (Sequential)")
                        stream_container.markdown(f"### {icon} {name}\nกำลังเชื่อมต่อ...")
                        full_ans = stream_container.write_stream(call_gemini_true_stream(multi_prompt, active_model, media_list=image_list if image_list else None, use_search=use_search))
                        if not full_ans or full_ans.startswith("❌"):
                            # 🔴 V9: Graceful skip — บันทึก error แต่ไม่หยุดระบบ
                            st.session_state.health_report[name] = f"❌ Failed — ข้ามและรันต่อ<br><span class='status-err'>{str(full_ans)[:100]}</span>"
                            st.session_state.raw_results.append((name, f"[⚠️ Agent นี้ไม่สำเร็จ — ข้ามไป]"))
                        else:
                            st.session_state.health_report[name] = "✅ Complete"
                            st.session_state.raw_results.append((name, full_ans))
                            st.session_state.meeting_log += f"### {icon} {name}\n{full_ans}\n\n"
                            # B-4: เพิ่มจาก 600 → 1000 chars ต่อ agent เพื่อไม่ให้ตัดกลางความคิด
                            meeting_context += f"--- {name} ---\n{str(full_ans)[:1000]}\n\n"

                    # รัน parallel batch
                    if parallel_ids:
                        status.update(label=f"⚡ กำลังรัน {len(parallel_ids)} Agent พร้อมกัน (Parallel Mode)...")
                        stream_container.markdown(f"⚡ **Parallel Processing:** {', '.join([team_data[aid]['icon']+team_data[aid]['name'] for aid in parallel_ids])}\n\n*(กรุณารอ ไม่มี live stream ในโหมดนี้)*")
                    
                        def run_agent(aid):
                            name = team_data[aid]['name']
                            icon = team_data[aid]['icon']
                            p = team_data[aid]['p']
                            prompt = build_prompt(aid, name, icon, p, meeting_context)
                            result = call_gemini_sync(prompt, active_model, media_list=image_list if image_list else None)
                            return (aid, name, icon, result)

                        # 🔴 V9: max_workers=3 (semaphore จำกัด concurrent calls)
                        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
                            futures = {executor.submit(run_agent, aid): aid for aid in parallel_ids}
                            for future in concurrent.futures.as_completed(futures):
                                aid, name, icon, full_ans = future.result()
                                if not full_ans or full_ans.startswith("❌"):
                                    # 🔴 V9: Graceful skip — บันทึกแต่ไม่หยุด
                                    st.session_state.health_report[name] = f"❌ Failed — ข้ามไป"
                                    st.session_state.raw_results.append((name, "[⚠️ Agent ไม่สำเร็จ — ข้ามไป]"))
                                else:
                                    st.session_state.health_report[name] = "✅ Complete"
                                    st.session_state.raw_results.append((name, full_ans))
                                    st.session_state.meeting_log += f"### {icon} {name}\n{full_ans}\n\n"

                else:
                    # ── Sequential mode (เดิม) ──
                    for idx, aid in enumerate(selected_agents):
                        name, icon, p = team_data[aid]['name'], team_data[aid]['icon'], team_data[aid]['p']

                        # ── B-3a: Compress context ทุก 3 agent ──────────────────────────
                        if idx > 0 and idx % 3 == 0:
                            status.update(label="👩‍💼 เลขานุการ กำลังย่อข้อมูล...")
                            compress_prompt = (
                                f"ย่อความเห็นของทีมต่อไปนี้ให้กระชับที่สุด "
                                f"โดยเก็บ key points สำคัญไว้ทั้งหมด ห้ามตัดข้อมูลสำคัญ:\n{meeting_context}"
                            )
                            sum_ans = "".join(list(call_gemini_true_stream(
                                compress_prompt, active_model, max_output_tokens=2048)))
                            meeting_context = f"--- 📋 สรุปความเห็น (รอบที่ {idx//3}) ---\n{sum_ans}\n\n"
                            st.session_state.meeting_log += f"### 👩‍💼 เลขานุการ (สรุป)\n> {sum_ans}\n\n"

                        # ── B-3b: Hard cap context ถ้ายาวเกิน 10,000 chars ──────────────
                        if len(meeting_context) > 10000:
                            # ตัดส่วนต้นออก เก็บแค่ส่วนหลัง (context ใหม่กว่าสำคัญกว่า)
                            meeting_context = "[...ตัดบางส่วน...]\n" + meeting_context[-8000:]

                        use_search = (aid == "A17")
                        multi_prompt = build_prompt(aid, name, icon, p, meeting_context, use_search)
                        prompt_size = len(multi_prompt)
                        status.update(label=f"{icon} แผนก{name} กำลังวิเคราะห์งาน... (prompt {prompt_size:,} chars)")
                        stream_container.markdown(f"### {icon} {name}\nกำลังเชื่อมต่อ Neural Engine...")
                        full_ans = stream_container.write_stream(
                            call_gemini_true_stream(multi_prompt, active_model,
                                media_list=image_list if image_list else None,
                                use_search=use_search)  # max_output_tokens=None → auto จาก B-1
                        )
                        if not full_ans or full_ans.startswith("❌"):
                            # 🔴 V9: Graceful skip
                            st.session_state.health_report[name] = f"❌ Failed — ข้ามไป<br><span class='status-err'>{str(full_ans)[:100]}</span>"
                            st.session_state.raw_results.append((name, "[⚠️ Agent ไม่สำเร็จ — ข้ามไป]"))
                        else:
                            st.session_state.health_report[name] = "✅ Complete"
                            st.session_state.raw_results.append((name, full_ans))
                            st.session_state.meeting_log += f"### {icon} {name}\n{full_ans}\n\n"
                            # B-5: เพิ่มจาก 600 → 1000 chars (parallel sequential phase)
                            meeting_context += f"--- {name} ---\n{str(full_ans)[:1000]}\n\n"

            except Exception as _run_err:
                _run_error = _run_err
                st.error(f"❌ เกิดข้อผิดพลาดระหว่างประชุม: {_run_err}")
            finally:
                # BUG1: ปลดล็อกเสมอ ไม่ว่าจะ success หรือ error
                st.session_state.is_running = False
            stream_container.empty()
            orbit_place.empty()
            status.update(label="✅ ประชุมเสร็จสิ้น!", state="complete")


            # FIX #9: บันทึก Meeting History ลง Vault
            timestamp = datetime.now().strftime("%d/%m/%Y %H:%M")
            history_entry = {
                "timestamp": timestamp,
                "brief": prompt_input[:120] + ("..." if len(prompt_input) > 120 else ""),
                "agents": len(selected_agents),
                "log": st.session_state.meeting_log
            }
            if "history" not in st.session_state.vault[st.session_state.current_project]:
                st.session_state.vault[st.session_state.current_project]["history"] = []
            st.session_state.vault[st.session_state.current_project]["history"].append(history_entry)
            st.session_state.vault[st.session_state.current_project]["history"] = \
                st.session_state.vault[st.session_state.current_project]["history"][-10:]
            save_vault(st.session_state.vault)

            # ── 🧠 MEMORY SYSTEM: Auto-generate memory summary หลังประชุมจบ ──
            if st.session_state.meeting_log and len(st.session_state.meeting_log) > 500:
                try:
                    with st.spinner("🧠 กำลังสรุปความรู้สะสมใหม่ให้ project..."):
                        _mem_model = get_best_model(API_KEY)
                        _new_memory = generate_memory_summary(
                            st.session_state.meeting_log, prompt_input, _mem_model
                        )
                        if _new_memory and not _new_memory.startswith("❌"):
                            save_project_memory(st.session_state.current_project, _new_memory)
                            st.session_state["memory_just_updated"] = True
                except Exception as _mem_err:
                    pass  # Memory generation ล้มเหลวไม่กระทบระบบหลัก

            # 💾 V9: บันทึก Cache
            if st.session_state.meeting_log and len(st.session_state.raw_results) > 0:
                st.session_state.brief_cache[session_cache_key] = {
                    "log": st.session_state.meeting_log,
                    "results": [[n, a] for n, a in st.session_state.raw_results],
                    "summary": st.session_state.summary_text,
                    "ts": timestamp
                }
                # เก็บแค่ 30 cache entries ล่าสุด
                if len(st.session_state.brief_cache) > 30:
                    oldest = list(st.session_state.brief_cache.keys())[0]
                    del st.session_state.brief_cache[oldest]
                save_cache(st.session_state.brief_cache)

            # 📊 V9: บันทึก Analytics
            record_session_analytics(
                st.session_state.current_project,
                selected_agents,
                st.session_state.session_token_count,
                st.session_state.session_cost_usd,
                st.session_state.ui_lang,
                brief=prompt_input,
            )
# ==========================================
    # FIX #9: MEETING HISTORY
    # ==========================================
    project_history = st.session_state.vault[st.session_state.current_project].get("history", [])
    if project_history:
        st.markdown("---")
        with st.expander(f"🕰️ ประวัติการประชุม ({len(project_history)} session ล่าสุด)", expanded=False):
            for i, entry in enumerate(reversed(project_history)):
                st.markdown(f"""<div class="history-card">
                    <b>📅 {entry['timestamp']}</b> · {entry['agents']} Agent<br>
                    <span style="color:#94a3b8;font-size:13px;">{entry['brief']}</span>
                </div>""", unsafe_allow_html=True)
                if st.button(f"📂 โหลดการประชุมนี้", key=f"load_hist_{i}"):
                    st.session_state.meeting_log = entry["log"]
                    st.rerun()


    # ==========================================
    # โซนผลลัพธ์
    # ==========================================
    if st.session_state.meeting_log:
        st.markdown("---")
        st.subheader("📌 บันทึกความเห็นจาก SPECIAL TEAM")

        # FIX #7: Agent Rating System
        if st.session_state.raw_results:
            with st.expander("⭐ ให้คะแนนแต่ละแผนก (Agent Rating)", expanded=False):
                rating_cols = st.columns(3)
                for i, (name, _) in enumerate(st.session_state.raw_results):
                    with rating_cols[i % 3]:
                        rating = st.select_slider(
                            f"{name}",
                            options=[1, 2, 3, 4, 5],
                            value=st.session_state.agent_ratings.get(name, 3),
                            key=f"rate_{name}"
                        )
                        st.session_state.agent_ratings[name] = rating
                        stars = "⭐" * rating + "☆" * (5 - rating)
                        st.caption(stars)

        # ✅ FIX: แสดงผลแบบ scrollable + แยก agent ชัดเจน
        log_text = st.session_state.meeting_log
        sections = log_text.split("\n---\n")
        for section in sections:
            section = section.strip()
            if not section:
                continue
            lines = section.split("\n")
            title_line = next((l for l in lines if l.startswith("###")), None)
            if title_line:
                title = title_line.replace("###", "").strip()
                body = "\n".join(l for l in lines if l != title_line).strip()
                with st.expander(title, expanded=True):
                    st.markdown(body)
            else:
                st.markdown(section)

        # FIX #10: DEBATE MODE — Round 2
        if debate_mode and st.session_state.raw_results:
            st.markdown("---")
            st.markdown('<div class="debate-badge">🥊 DEBATE MODE — Round 2: Agent ถกเถียงกัน</div>', unsafe_allow_html=True)
            if st.button("🔥 เริ่ม Debate Round 2", use_container_width=True):
                active_model = get_best_model(API_KEY)
                debate_status = st.status("🥊 กำลังให้ Agent comment งานกันและกัน...")
                debate_log = ""
                # ให้แต่ละ agent อ่านงานคนอื่นแล้ว comment/ท้วง
                for aid in selected_agents[:25]:  # B-7: รองรับทุก agent
                    name, icon, p = team_data[aid]['name'], team_data[aid]['icon'], team_data[aid]['p']
                    debate_prompt = f"""คุณคือ {name} ({p})
# BUG4: 5000 chars ไม่พอสำหรับ 25 agents → เพิ่มเป็น 12000
นี่คือสรุปความเห็นของทีมในรอบแรก:
{st.session_state.meeting_log[:12000]}

งานของคุณตอนนี้: อ่านความเห็นของเพื่อนร่วมทีม แล้ว:
1. ชี้ให้เห็น จุดที่คุณ **เห็นด้วย** พร้อมเหตุผล
2. ชี้ให้เห็น จุดที่คุณ **ไม่เห็นด้วย หรือมีข้อท้วง** พร้อมเสนอทางเลือกที่ดีกว่า
3. เพิ่มเติมข้อมูลที่ทีมยังขาด จากมุมมองของ{name}โดยเฉพาะ"""
                    debate_status.update(label=f"{icon} {name} กำลัง debate...")
                    debate_ans = "".join(list(call_gemini_true_stream(debate_prompt, active_model)))
                    debate_log += f"### 🥊 {icon} {name} (Debate)\n{debate_ans}\n\n"
                    st.session_state.health_report[f"{name} (Debate)"] = "✅ Complete"
                st.session_state.meeting_log += f"\n---\n## 🥊 Debate Round 2\n{debate_log}"
                debate_status.update(label="✅ Debate เสร็จสิ้น!", state="complete")
                st.rerun()

        # ── Chat กับทีม ──
        st.markdown("#### 💬 สื่อสารกับทีม AQUALINE เพิ่มเติม")
        user_chat = st.text_input("พิมพ์ข้อความเพื่อสอบถาม, แก้ไข หรือสั่งการทีมเพิ่มเติม:")
        if st.button("📤 ส่งข้อความถึงทีม"):
            if user_chat:
                with st.spinner("ทีมกำลังตอบกลับ..."):
                    active_model = get_best_model(API_KEY)
                    prompt = f"นี่คือบันทึกการประชุมที่ผ่านมาทั้งหมด:\n{st.session_state.meeting_log}\n\nคำสั่งจากหัวหน้าทีม: {user_chat}\nตอบรับและนำเสนอแนวทาง{lang_suffix(st.session_state.ui_lang)}"
                    team_reply = "".join(list(call_gemini_true_stream(prompt, active_model)))
                    st.session_state.meeting_log += f"---\n\n### 👤 หัวหน้าทีม\n**{user_chat}**\n\n### 🤖 ทีม SPECIAL TEAM\n{team_reply}\n\n"
                    st.rerun()

        # 🧮 V9: Token Counter Display
        if st.session_state.session_token_count > 0:
            tok = st.session_state.session_token_count
            cost_thb = st.session_state.session_cost_usd * 36  # approx THB
            st.markdown(f"""
            <div class="token-card">
                <div style="display:flex;gap:30px;flex-wrap:wrap;">
                    <div><div class="token-num">{tok:,}</div><div class="token-label">Tokens Used (est.)</div></div>
                    <div><div class="token-num">${st.session_state.session_cost_usd:.4f}</div><div class="token-label">Cost USD (est.)</div></div>
                    <div><div class="token-num">฿{cost_thb:.2f}</div><div class="token-label">Cost THB (est.)</div></div>
                </div>
                <div style="font-size:10px;color:#6ee7b7;margin-top:6px;">* ประมาณการจาก Gemini 2.5 Flash pricing · 1 token ≈ 4 chars</div>
            </div>""", unsafe_allow_html=True)

        # ── Export ──
        st.markdown("#### 📦 Export ผลลัพธ์")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            if st.button("📊 สรุป Action Plan ตาราง", use_container_width=True):
                with st.spinner("กำลังสรุป..."):
                    st.session_state.summary_text = "".join(list(call_gemini_true_stream(
                        f"สรุปการประชุมนี้เป็นตาราง Action Plan:\n{st.session_state.meeting_log[:10000]}{lang_suffix(st.session_state.ui_lang)}",
                        get_best_model(API_KEY)
                    )))
        with col2:
            if st.button("📝 สรุปเนื้อความเชิงกลยุทธ์", use_container_width=True):
                with st.spinner("กำลังสรุป..."):
                    st.session_state.summary_text = "".join(list(call_gemini_true_stream(
                        f"สรุปการประชุมนี้เป็นเนื้อความเชิงกลยุทธ์:\n{st.session_state.meeting_log[:10000]}{lang_suffix(st.session_state.ui_lang)}",
                        get_best_model(API_KEY)
                    )))
        with col3:
            if not st.session_state.ready_to_export:
                if st.button("📄 เตรียม Word (.docx)", use_container_width=True):
                    st.session_state.ready_to_export = True; st.rerun()
            else:
                doc = Document()
                doc.add_heading("AQUALINE REPORT — V9.0 ULTRA", 0)
                doc.add_heading("รายละเอียดบรีฟ", 1)
                # BUG7: ถ้า prompt_input หายหลัง rerun ให้ fallback จาก vault
                _brief_for_export = prompt_input or current_data.get("brief", "")
                doc.add_paragraph(_brief_for_export)
                for name, ans in st.session_state.raw_results:
                    doc.add_heading(f"แผนก: {name}", 2)
                    doc.add_paragraph(ans)
                if st.session_state.summary_text:
                    doc.add_heading('Action Plan', 1)
                    doc.add_paragraph(st.session_state.summary_text)
                if st.session_state.agent_ratings:
                    doc.add_heading('Agent Ratings', 1)
                    for rname, rval in st.session_state.agent_ratings.items():
                        doc.add_paragraph(f"{rname}: {'⭐'*rval}")
                bio = BytesIO()
                doc.save(bio)
                st.download_button("📥 ดาวน์โหลด Word", data=bio.getvalue(),
                                   file_name=f"Report_{st.session_state.current_project}.docx", use_container_width=True)

        # 📄 V9: PDF Export
        with col4:
            if st.button("🎨 Export PDF (Aqualine Brand)", use_container_width=True):
                if REPORTLAB_OK:
                    with st.spinner("🖨️ กำลังสร้าง PDF..."):
                        pdf_bytes = generate_branded_pdf(
                            st.session_state.current_project,
                            prompt_input,
                            st.session_state.raw_results,
                            st.session_state.summary_text,
                            st.session_state.agent_ratings
                        )
                    if pdf_bytes:
                        st.download_button(
                            "📥 ดาวน์โหลด PDF",
                            data=pdf_bytes,
                            file_name=f"Aqualine_Report_{st.session_state.current_project}_{datetime.now().strftime('%Y%m%d')}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                else:
                    st.warning("⚠️ กรุณาติดตั้ง reportlab ก่อน: `pip install reportlab`")

        if st.session_state.summary_text:
            st.info("📑 ผลสรุปงานล่าสุด")
            st.markdown(st.session_state.summary_text)

        # 🔍 V9: AI Keyword Extractor
        st.markdown("---")
        st.markdown("#### 🔍 AI Keyword Extractor — สกัด Keywords, Actions, Deadlines")
        if "extracted_keywords" not in st.session_state: st.session_state.extracted_keywords = None
        if st.button("🤖 ให้ AI สกัด Keywords & Action Items อัตโนมัติ", use_container_width=True):
            with st.spinner("🔍 AI กำลังวิเคราะห์ผลการประชุม..."):
                active_model = get_best_model(API_KEY)
                st.session_state.extracted_keywords = extract_keywords_and_actions(
                    st.session_state.meeting_log, active_model, st.session_state.ui_lang
                )
        if st.session_state.extracted_keywords:
            kw_data = st.session_state.extracted_keywords
            kw_cols = st.columns(2)
            with kw_cols[0]:
                st.markdown("**🏷️ Keywords สำคัญ**")
                kw_html = "".join([f'<span class="kw-chip">{k}</span>' for k in kw_data.get("keywords", [])])
                st.markdown(kw_html or "ไม่พบ", unsafe_allow_html=True)
                st.markdown("**⚠️ ความเสี่ยง**")
                for r in kw_data.get("risks", []):
                    st.markdown(f"- {r}")
            with kw_cols[1]:
                st.markdown("**✅ Action Items**")
                action_html = "".join([f'<span class="action-chip">{a}</span>' for a in kw_data.get("action_items", [])])
                st.markdown(action_html or "ไม่พบ", unsafe_allow_html=True)
                st.markdown("**📅 Deadlines / Timeline**")
                for d in kw_data.get("deadlines", []):
                    st.markdown(f"- {d}")
            if kw_data.get("key_decisions"):
                st.markdown("**🎯 การตัดสินใจสำคัญ**")
                for kd in kw_data["key_decisions"]:
                    st.markdown(f"- {kd}")


    # ==========================================
    # 👑 CHAIRMAN MASTER
    # ==========================================
    st.markdown("---")
    st.markdown("""
    <div style='background:linear-gradient(90deg,rgba(83,74,183,.15),rgba(219,39,119,.1));
    border:1px solid #534AB7;border-radius:12px;padding:16px 20px;margin-bottom:4px'>
    <span style='font-size:22px;font-weight:900;color:#fff'>👑 CHAIRMAN MASTER</span>
    <span style='font-size:12px;color:#a78bfa;margin-left:10px'>
    ประธานสูงสุด — สรุปและวิเคราะห์ภาพรวมจากทุก Agent
    </span>
    </div>
    """, unsafe_allow_html=True)

    # ── AI Mode Selector ──
    st.markdown("**🤖 เลือก AI สำหรับ Chairman:**")
    _ai_mode_cols = st.columns(2)
    with _ai_mode_cols[0]:
        _gemini_selected = st.checkbox(
            "🟢 Google Gemini 2.5 Pro  (ฟรี)",
            value=(st.session_state.chairman_ai_mode == "gemini"),
            key="chk_gemini_chairman"
        )
    with _ai_mode_cols[1]:
        _claude_selected = st.checkbox(
            "🟣 Claude Haiku  (จำกัด 175 ฿/เดือน)",
            value=(st.session_state.chairman_ai_mode == "claude"),
            key="chk_claude_chairman"
        )

    if _gemini_selected and st.session_state.chairman_ai_mode != "gemini":
        st.session_state.chairman_ai_mode = "gemini"; st.rerun()
    if _claude_selected and st.session_state.chairman_ai_mode != "claude":
        st.session_state.chairman_ai_mode = "claude"; st.rerun()
    if not _gemini_selected and not _claude_selected:
        st.session_state.chairman_ai_mode = "gemini"; st.rerun()

    _mode_now = st.session_state.chairman_ai_mode

    # ── Info bar ──
    if _mode_now == "gemini":
        st.markdown("""
        <div style='background:rgba(16,185,129,.08);border:1px solid #10b981;
        border-radius:8px;padding:8px 14px;margin:6px 0 10px;font-size:12px;color:#6ee7b7'>
        ✅ <b>Gemini mode</b> — ใช้ quota Google AI Studio ของคุณ ไม่เสียเงินเพิ่ม
        &nbsp;|&nbsp; model: <b>Gemini 2.5 Pro</b>
        </div>
        """, unsafe_allow_html=True)
    else:
        _bd       = load_chairman_budget()
        _used_thb = _bd.get("used_thb", 0.0)
        _left_thb = max(0.0, CHAIRMAN_BUDGET_THB - _used_thb)
        _pct      = min(100, (_used_thb / CHAIRMAN_BUDGET_THB) * 100)
        _reset_date = (datetime.now().replace(day=1) + timedelta(days=32)).replace(day=1).strftime("%d/%m/%Y")
        _bar_color  = "#ef4444" if _pct >= 100 else "#f59e0b" if _pct >= 80 else "#8b5cf6"
        if _pct >= 100:
            _badge = f"<span style='background:#501313;color:#f7c1c1;padding:2px 8px;border-radius:20px;font-size:11px'>🔴 หมดแล้ว reset {_reset_date}</span>"
        elif _pct >= 80:
            _badge = f"<span style='background:#412402;color:#fac775;padding:2px 8px;border-radius:20px;font-size:11px'>🟡 เหลือน้อย {_left_thb:.2f} ฿</span>"
        else:
            _badge = f"<span style='background:#26215C;color:#c4b5fd;padding:2px 8px;border-radius:20px;font-size:11px'>🟢 เหลือ {_left_thb:.2f} ฿</span>"
        st.markdown(f"""
        <div style='background:rgba(83,74,183,.08);border:1px solid #534AB7;
        border-radius:8px;padding:8px 14px;margin:6px 0 10px'>
        <div style='display:flex;justify-content:space-between;align-items:center;margin-bottom:5px'>
            <span style='font-size:12px;color:#94a3b8'>
            🟣 <b>Claude Haiku</b> &nbsp;|&nbsp; งบเดือนนี้ {_used_thb:.2f} / {CHAIRMAN_BUDGET_THB:.0f} ฿
            </span>{_badge}
        </div>
        <div style='background:#1e293b;border-radius:4px;height:6px;overflow:hidden'>
            <div style='background:{_bar_color};width:{_pct:.1f}%;height:100%;border-radius:4px'></div>
        </div>
        </div>
        """, unsafe_allow_html=True)

    # ── ผลเดิม (ถ้ามี) ──
    if st.session_state.chairman_result:
        st.markdown("""
        <div style='background:rgba(83,74,183,.06);border:1px solid #534AB7;
        border-radius:10px;padding:14px 18px;margin-bottom:10px'>
        <div style='font-size:13px;font-weight:600;color:#a78bfa;margin-bottom:8px'>
        👑 ผลการวิเคราะห์ของ Chairman Master
        </div>
        """, unsafe_allow_html=True)
        st.markdown(st.session_state.chairman_result)
        if st.session_state.chairman_cost_thb > 0:
            st.markdown(f"""
            <div style='margin-top:8px;font-size:12px;color:#6ee7b7;
            background:rgba(16,185,129,.1);border-radius:6px;padding:6px 12px'>
            💸 ครั้งนี้เสียไป <b>{st.session_state.chairman_cost_thb:.4f} ฿</b> (Claude Haiku)
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div style='margin-top:8px;font-size:12px;color:#6ee7b7;
            background:rgba(16,185,129,.08);border-radius:6px;padding:6px 12px'>
            ✅ ใช้ Gemini — ไม่เสียเงินเพิ่ม
            </div>
            """, unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

    # ── Input + Button ──
    if not st.session_state.meeting_log:
        st.info("⏳ ต้องประชุมก่อน — กด 'เริ่มระดมสมอง' แล้วค่อยกด Chairman")
    elif _mode_now == "claude" and get_chairman_budget_used() >= CHAIRMAN_BUDGET_THB:
        _reset_date = (datetime.now().replace(day=1) + timedelta(days=32)).replace(day=1).strftime("%d/%m/%Y")
        st.error(f"🔴 งบ Claude หมดแล้ว reset {_reset_date} — หรือสลับไปใช้ Gemini แทนได้เลย")
    else:
        _chairman_brief = st.text_area(
            "📝 คำสั่งพิเศษถึง Chairman (ไม่บังคับ — ถ้าว่างจะสรุปแบบ full detail อัตโนมัติ):",
            placeholder="เช่น: สรุปเฉพาะขั้นตอนยิงแอด Facebook\nหรือ: วิเคราะห์ว่าภาพนี้เหมาะกลุ่ม Gen Z มั้ย",
            height=80, key="chairman_custom_brief"
        )
        if _mode_now == "claude":
            _pct_now = min(100, (get_chairman_budget_used() / CHAIRMAN_BUDGET_THB) * 100)
            if _pct_now >= 80:
                _left_now = max(0, CHAIRMAN_BUDGET_THB - get_chairman_budget_used())
                st.warning(f"⚠️ งบ Claude เหลือแค่ {_left_now:.2f} ฿")

        _btn_label = (
            "👑 เรียก CHAIRMAN — Gemini 2.5 Pro (ฟรี)"
            if _mode_now == "gemini"
            else "👑 เรียก CHAIRMAN — Claude Haiku (เสียเงิน)"
        )

        if st.button(_btn_label, use_container_width=True, key="btn_chairman"):
            _current_brief = st.session_state.vault.get(
                st.session_state.current_project, {}
            ).get("brief", "")
            _sys_p, _usr_p = build_chairman_prompt(
                brief=_current_brief,
                meeting_log=st.session_state.meeting_log,
                custom_instruction=_chairman_brief,
                lang=st.session_state.ui_lang
            )
            _container = st.empty()
            _container.markdown("👑 **Chairman กำลังวิเคราะห์...**")

            if _mode_now == "gemini":
                _pro_model = get_best_model(API_KEY)
                try:
                    _r = requests.get(
                        f"https://generativelanguage.googleapis.com/v1beta/models?key={API_KEY}",
                        timeout=5)
                    _avail = [m["name"] for m in _r.json().get("models", [])
                              if "generateContent" in m.get("supportedGenerationMethods", [])]
                    for _m in ["models/gemini-2.5-pro",
                               "models/gemini-2.5-flash-preview-05-20",
                               "models/gemini-2.5-flash"]:
                        if _m in _avail:
                            _pro_model = _m; break
                except: pass

                with st.spinner(f"👑 Chairman (Gemini) กำลังวิเคราะห์..."):
                    _full_answer = _container.write_stream(
                        call_gemini_chairman(_sys_p, _usr_p, _pro_model)
                    )
                st.session_state.chairman_result   = str(_full_answer)
                st.session_state.chairman_cost_thb = 0.0
                st.success("✅ Chairman (Gemini) เสร็จ! — ไม่เสียเงิน ✨")
            else:
                if not ANTHROPIC_API_KEY:
                    st.error("❌ ไม่พบ ANTHROPIC_API_KEY\nเพิ่มใน .streamlit/secrets.toml")
                else:
                    with st.spinner("👑 Chairman (Claude Haiku) กำลังวิเคราะห์..."):
                        _full_answer = _container.write_stream(
                            call_claude_chairman(_sys_p, _usr_p)
                        )
                    _cost = add_chairman_cost_claude(_usr_p, str(_full_answer))
                    st.session_state.chairman_result   = str(_full_answer)
                    st.session_state.chairman_cost_thb = _cost
                    st.success(f"✅ Chairman (Claude) เสร็จ! เสียไป {_cost:.4f} ฿")

            _ai_label = "Gemini 2.5 Pro" if _mode_now == "gemini" else "Claude Haiku"
            st.session_state.meeting_log += (
                f"\n---\n## 👑 CHAIRMAN MASTER ({_ai_label})\n"
                f"{st.session_state.chairman_result}\n\n"
            )
            st.rerun()

    # ==========================================
    # 🌌 Andromeda Pre-Meeting Auditor
    # ==========================================
    st.markdown("---")
    st.markdown("### 🌌 Andromeda Pre-Meeting Auditor (Omni Scanner)")
    
    audit_files = st.file_uploader("📥 อัปโหลดไฟล์ภาพ/เอกสาร/วิดีโอ (mp4) ให้ Andromeda ตรวจสอบ",
                                    type=["jpg","jpeg","png","pdf","docx","mp4","mov"],
                                    accept_multiple_files=True, key="audit_uploader")
    
    if st.button("🔍 ส่งให้ Andromeda สแกนและสรุปกฎเหล็ก"):
        if not audit_files:
            st.error("⚠️ กรุณาแนบไฟล์ที่ต้องการให้ Andromeda สแกนก่อนครับ")
        else:
            with st.spinner("🌌 Andromeda กำลังสแกน..."):
                audit_text = ""
                media_for_andromeda = None
                for file in audit_files:
                    if file.type == "application/pdf":
                        with pdfplumber.open(file) as pdf:
                            audit_text += "\n".join([p.extract_text() or "" for p in pdf.pages])
                    elif "wordprocessingml" in file.type:
                        doc = Document(file)
                        audit_text += "\n".join([p.text for p in doc.paragraphs])
                    elif "video" in file.type or "image" in file.type:
                        if not media_for_andromeda:
                            media_for_andromeda = {"type": file.type, "b64": base64.b64encode(file.read()).decode('utf-8')}
                        audit_text += f"\n[ไฟล์สื่อ {file.name} ถูกแนบ]"
                audit_prompt = f"ไฟล์บรีฟ/สคริปต์:\n{audit_text}\n\nจงสรุปเป็น 'ข้อควรระวังและกฎเหล็ก' ให้ทีม Special Team นำไปใช้"
                sys_instruct = "คุณคือ Andromeda (Ads Optimizer) กฎเหล็ก: 1. เช็ค Hook Rate 3 วิแรก 2. ตรวจสอบ Policy โฆษณา 3. ให้คะแนนความน่าดึงดูดใจ"
                st.session_state.audit_result = "".join(list(call_gemini_true_stream(
                    audit_prompt, get_best_model(API_KEY),
                    media_data=media_for_andromeda, system_instruction=sys_instruct
                )))
                st.success("✅ Andromeda สแกนเรียบร้อย!")
                
    if "audit_result" in st.session_state and st.session_state.audit_result:
        st.markdown("<div style='background:#0b1426;color:#4A90E2;padding:15px;border:1px solid #0668E1;border-radius:8px;'>", unsafe_allow_html=True)
        st.markdown("#### 📜 กฎเหล็กจากเรดาร์ Andromeda:")
        st.markdown(st.session_state.audit_result)
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("---")

# ── Health Report ──
with col_health:
    st.subheader("🩺 Health Report")

    # 🔍 Model Debug
    with st.expander("🔮 ตรวจสอบโมเดล API", expanded=True):
        selected_model = get_best_model(API_KEY)
        st.success(f"✅ โมเดลที่ใช้: `{selected_model}`")
        if st.button("🔄 รีเฟรชโมเดล", use_container_width=True):
            get_best_model.clear()
            if "cached_model_list" in st.session_state:
                del st.session_state.cached_model_list  # BUG8: ล้าง cache ด้วย
            st.rerun()
        # BUG8: ไม่ควร call API ทุก rerender → cache ใน session_state
        if "cached_model_list" not in st.session_state:
            try:
                _res = requests.get(f"https://generativelanguage.googleapis.com/v1beta/models?key={API_KEY}", timeout=8)
                if _res.status_code == 200:
                    st.session_state.cached_model_list = [
                        m["name"] for m in _res.json().get("models",[])
                        if "generateContent" in m.get("supportedGenerationMethods",[])
                    ]
                else:
                    st.session_state.cached_model_list = []
            except:
                st.session_state.cached_model_list = []
        all_models = st.session_state.cached_model_list
        if all_models:
            st.caption(f"พบ {len(all_models)} โมเดล:")
            for m in all_models:
                badge = "🟢 ใช้อยู่" if m == selected_model else "⚪"
                st.caption(f"{badge} {m}")
        else:
            st.caption("ไม่สามารถดึงรายการโมเดลได้")

    st.markdown("<div style='font-size:12px;color:#94a3b8;margin-bottom:15px;'>เช็คสถานะ Error ของแต่ละแผนก</div>", unsafe_allow_html=True)
    if not st.session_state.health_report:
        st.write("ยังไม่ได้เริ่มรันระบบ")
    else:
        for name, status_txt in st.session_state.health_report.items():
            st.markdown(f"""<div class="status-card"><b>{name}</b><br>{status_txt}</div>""", unsafe_allow_html=True)
    
    # แสดง Agent Ratings ใน Health Column ด้วย
    if st.session_state.agent_ratings:
        st.markdown("---")
        st.markdown("**⭐ คะแนน Agent**")
        for rname, rval in st.session_state.agent_ratings.items():
            stars = "⭐" * rval
            st.markdown(f"<div class='status-card'><b>{rname}</b><br>{stars}</div>", unsafe_allow_html=True)

    # 📊 V9: Analytics Dashboard
    st.markdown("---")
    with st.expander("📊 Analytics Dashboard", expanded=False):
        a = st.session_state.analytics
        total_sessions = len(a.get("sessions", []))
        total_tokens = sum(s.get("tokens", 0) for s in a.get("sessions", []))
        total_cost = sum(s.get("cost_usd", 0) for s in a.get("sessions", []))

        st.markdown(f"""
        <div class="analytics-card">
            <div class="analytics-num">{total_sessions}</div>
            <div class="analytics-label">TOTAL SESSIONS</div>
        </div>
        <div class="analytics-card">
            <div class="analytics-num">{total_tokens:,}</div>
            <div class="analytics-label">TOTAL TOKENS (est.)</div>
        </div>
        <div class="analytics-card">
            <div class="analytics-num">${total_cost:.3f}</div>
            <div class="analytics-label">TOTAL COST USD (est.)</div>
        </div>
        """, unsafe_allow_html=True)

        # Top Agents
        if a.get("agent_usage"):
            st.markdown("**🏆 Agent ที่ใช้บ่อยที่สุด**")
            sorted_agents = sorted(a["agent_usage"].items(), key=lambda x: x[1], reverse=True)[:5]
            for aid, cnt in sorted_agents:
                agent_name = team_data.get(aid, {}).get("name", aid)
                agent_icon = team_data.get(aid, {}).get("icon", "🤖")
                st.markdown(f"<div class='status-card'>{agent_icon} {agent_name}<br><b>{cnt} ครั้ง</b></div>", unsafe_allow_html=True)

        # Top Projects
        if a.get("project_count"):
            st.markdown("**📁 Project ที่ประชุมบ่อย**")
            sorted_proj = sorted(a["project_count"].items(), key=lambda x: x[1], reverse=True)[:3]
            for pname, pcnt in sorted_proj:
                st.markdown(f"<div class='status-card'>📁 {pname}<br><b>{pcnt} session</b></div>", unsafe_allow_html=True)

        # Weekly trend
        if a.get("weekly"):
            st.markdown("**📅 การใช้งานรายสัปดาห์ (5 สัปดาห์ล่าสุด)**")
            sorted_weeks = sorted(a["weekly"].items())[-5:]
            for wk, wcount in sorted_weeks:
                bar = "█" * min(wcount, 20)
                st.caption(f"{wk}: {bar} ({wcount})")

        if st.button("🗑️ ล้างข้อมูล Analytics", use_container_width=True):
            st.session_state.analytics = {"sessions": [], "agent_usage": {}, "project_count": {}, "weekly": {}}
            save_analytics(st.session_state.analytics)
            st.success("ล้างแล้ว!")
            st.rerun()